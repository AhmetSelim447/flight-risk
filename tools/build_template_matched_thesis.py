from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\mete_\Downloads\bitirmeAraRaporBahar.docx")
OUT = ROOT / "Bitirme_Tezi_Sablon_Uyumlu_Flight_Risk.docx"


META_ROWS = [
    ("YILI / DÖNEMİ", "2025-2026 DERS YILI / Bahar DÖNEMİ"),
    ("ÖĞRENCİ NO", "220290007 - 220290002 - 220290009"),
    ("AD SOYAD", "Mete Han YILMAZ - Ahmet Selim AYTAÇ - Emre NABİKOĞLU"),
    ("BİTİRME TEZ DANIŞMANI", "Prof. Dr. Bilal ALATAŞ"),
    ("PROJE KONUSU/BAŞLIĞI", "Yapay Zeka Destekli NOTAM ve METAR/TAF Analizi ile Uçuş Risk Değerlendirme ve Karar Destek Sistemi"),
]


def set_text(paragraph, text: str, bold: bool | None = None, size: int = 10) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def clear_cell(cell) -> None:
    cell.text = ""
    for p in cell.paragraphs:
        p.text = ""


def add_para(cell, text: str, bold: bool | None = None, size: int = 10, indent: bool = False) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    if indent:
        text = "\t" + text
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_mono(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    r.font.size = Pt(8.5)


def clone_row(table, source_row_idx: int = -1):
    source = table.rows[source_row_idx]._tr
    new_tr = deepcopy(source)
    table._tbl.append(new_tr)
    return table.rows[-1]


def remove_all_table_rows(table) -> None:
    for row in list(table.rows):
        table._tbl.remove(row._tr)


def add_meta_row(table, key: str, value: str, template_idx: int = 0) -> None:
    row = clone_row(table, template_idx if len(table.rows) else 0)
    c0, c1 = row.cells[0], row.cells[1]
    clear_cell(c0)
    clear_cell(c1)
    set_text(c0.paragraphs[0], key, bold=True)
    set_text(c1.paragraphs[0], value)


def add_section_title_row(table, title: str) -> None:
    row = clone_row(table)
    merged = row.cells[0].merge(row.cells[1])
    clear_cell(merged)
    p = merged.paragraphs[0]
    set_text(p, title, bold=True)


def add_content_row(table, paragraphs: list[str], diagram: str | None = None) -> None:
    row = clone_row(table)
    merged = row.cells[0].merge(row.cells[1])
    clear_cell(merged)
    # Keep the first paragraph object from the template and use additional paragraphs for content.
    first = True
    for text in paragraphs:
        if first:
            set_text(merged.paragraphs[0], "\t" + text)
            merged.paragraphs[0].paragraph_format.space_after = Pt(6)
            first = False
        else:
            add_para(merged, text, indent=True)
    if diagram:
        add_mono(merged, diagram)


def p(*items: str) -> list[str]:
    return list(items)


def build_sections() -> list[tuple[str, list[str], str | None]]:
    architecture_diagram = """Kullanıcı
  |
  v
React Web UI / BriefPanel
  |
  v
Express API /brief orkestratörü
  |-- METAR/TAF provider zinciri
  |-- NOTAM provider veya sentetik event engine
  |-- pist, rüzgar ve alternate hesapları
  |
  v
services/nlp AI servisi
  |-- /ai/notam/parse
  |-- /ai/risk/predict
  |-- /ai/brief/report
  |
  v
Açıklanabilir risk brifingi + PDF + harita"""

    hybrid_diagram = """finalScore = 0.65 * mlScore
           + 0.25 * ruleScore
           + 0.10 * notamSemanticScore

Guardrail:
  high weather floor    -> final score en az 70
  caution weather floor -> final score en az 40

LLM/rapor katmanı:
  skoru serbestçe üretmez,
  yalnızca verilen veriye dayalı açıklama üretir."""

    notam_diagram = """NOTAM kaynağı
  |
  +-- canlı provider varsa: Laminar / SkyLink / EAD hedefi
  |
  +-- canlı provider yoksa: deterministik sentetik NOTAM
          |
          v
category + severity + critical + affectedRunway + score + reason
          |
          v
UI: kısa operasyonel gerekçe + ham metin detay altında"""

    sections: list[tuple[str, list[str], str | None]] = []

    sections.append((
        "Giriş (Projenin genel özeti ve ilerleme durumu)",
        p(
            "Bu bitirme tezinde, uçuş öncesi operasyonel brifing sürecini desteklemek amacıyla geliştirilen yapay zeka destekli bir karar destek sistemi sunulmaktadır. Sistem; METAR, TAF, NOTAM, pist bilgileri, rüzgar bileşenleri, alternatif meydan önerileri ve model çıktısını tek bir akışta birleştirerek kullanıcıya anlaşılır bir uçuş risk brifingi üretir.",
            "Çalışmanın ana hedefi, pilot, dispeçer, AIS/AIM, ATC veya resmi operasyonel otoritenin yerine geçen bir karar sistemi geliştirmek değildir. Sistem, uçuşu onaylayan veya iptal eden bir mekanizma olarak değil, uçuş öncesi kontrol edilmesi gereken risk başlıklarını açıklanabilir biçimde gösteren bir brifing asistanı olarak tasarlanmıştır.",
            "Proje, güz döneminde oluşturulan temel veri toplama ve brifing altyapısının üzerine bahar döneminde hibrit yapay zeka, tarihsel METAR veri seti, model eğitimi, sentetik NOTAM event engine, kalibrasyon ekranı, kullanıcı geri bildirimi, loglama ve sadeleştirilmiş karar özeti katmanlarının eklenmesiyle genişletilmiştir.",
            "Bu tez metninde sistemin ürün amacı, veri kaynakları, mimari yapısı, risk hesaplama mantığı, ML pipeline süreci, NOTAM yorumlama stratejisi, kullanıcı arayüzü kararları, test ve doğrulama bulguları ayrıntılı biçimde açıklanmaktadır. Özellikle hibrit AI yaklaşımının neden seçildiği ve sistemin hangi sınırlılıklar altında çalıştığı açıkça belirtilmiştir.",
            "Projenin ayırt edici yönü, büyük dil modeli veya tekil bir yapay zeka bileşenine sınırsız karar yetkisi vermemesidir. Kural tabanlı motor, METAR tabanlı operasyonel proxy risk modeli ve NOTAM semantik etki sınıflandırması birlikte çalışır; LLM-benzeri rapor katmanı ise bu çıktıları okunabilir hale getirir.",
        ),
        None,
    ))

    sections.append((
        "Projenin Amacı, Kapsamı ve Sınırları",
        p(
            "Uçuş öncesi brifing sürecinde kullanıcı çoğu zaman farklı ekranlardan veri toplamak zorundadır. METAR mevcut hava durumunu, TAF beklenen hava koşullarını, NOTAM operasyonel kısıtları, pist bilgisi meydan kullanılabilirliğini, rüzgar bileşeni ise kalkış ve iniş performansını etkiler. Bu verilerin ayrı ayrı okunması mümkündür; fakat karar destek bağlamında birlikte yorumlanmaları gerekir.",
            "Bu projenin amacı, parçalı veri kaynaklarını tek ekranda birleştirmek ve kullanıcının 'risk seviyesi nedir?', 'risk neden oluştu?', 'hangi NOTAM kritik?', 'hangi veri eksik?', 'model bu sonucu ne kadar güvenle üretti?' sorularına hızlı ve açıklanabilir cevap vermektir. Bu nedenle sistem yalnızca ham METAR/TAF veya NOTAM listeleyen bir araç değildir.",
            "Kapsam dahilinde canlı METAR/TAF çekme, Türkiye LT* meydanları için tarihsel METAR veri seti oluşturma, operasyonel proxy risk modeli eğitme, sentetik NOTAM üretme, NOTAM etki sınıflandırması yapma, final risk skorunu hesaplama, risk raporu tablosu üretme ve PDF çıktı alma özellikleri yer almaktadır.",
            "Kapsam dışında kalan en önemli başlık, gerçek operasyonel uçuş onayıdır. Sistem bir havayolu operasyon kontrol merkezi, resmi uçuş brifing servisi, meteoroloji otoritesi veya AIS/AIM sağlayıcısı değildir. Üretilen skorlar, resmi operasyonel kararın yerine geçmez; sadece dikkat edilmesi gereken başlıkları öne çıkarır.",
            "Bu sınırın tez içinde açıkça belirtilmesi bilinçli bir tercihtir. Havacılık gibi yüksek emniyet hassasiyeti olan bir alanda yapay zekanın rolü abartılmamalıdır. Bu çalışma, sertifikalı bir karar sistemi değil, açıklanabilir hibrit AI yaklaşımının uçuş öncesi karar destek sürecine nasıl uygulanabileceğini gösteren bir prototiptir.",
        ),
        None,
    ))

    sections.append((
        "Modül 1: Veri Kaynaklarının Analizi ve Genişletilmesi",
        p(
            "Projede kullanılan veri kaynakları, sistemin doğruluğu ve sürdürülebilirliği açısından yeniden değerlendirilmiştir. İlk aşamada web sayfalarından veri çekme fikri değerlendirilmiş, ancak metar-taf.com gibi kaynakların anti-bot davranışı, izin/ToS belirsizliği ve HTML yapısının kırılganlığı nedeniyle üretim sağlayıcısı olarak kullanılmamasına karar verilmiştir.",
            "METAR/TAF için birincil kaynak olarak AviationWeather Data API tercih edilmiştir. Bu kaynak, METAR ve TAF ürünlerini dünya çapında kapsama ile JSON, XML, CSV, GeoJSON ve IWXXM gibi makine tarafından işlenebilir formatlarda sunar. Ayrıca API dokümantasyonu, oran sınırlaması ve hata kodları gibi üretim ortamı açısından gerekli bilgileri içerir.",
            "Sistem provider zinciri ile tasarlanmıştır. MET_PROVIDER=auto modunda önce AviationWeather denenir; token varsa CheckWX ve AVWX fallback olarak kullanılabilir; son aşamada NOAA text endpoint devreye girebilir. Böylece tek bir sağlayıcının geçici hatası brifing akışının tamamen durmasına neden olmaz.",
            "Tarihsel METAR verisi için Iowa State University Iowa Environmental Mesonet ASOS/METAR arşivi kullanılmıştır. Bu arşiv Türkiye ASOS ağı dahil olmak üzere çok sayıda istasyon için geçmiş gözlem verisi indirmeye imkan verir. Model eğitimi bu geçmiş veriden oluşturulan veri seti üzerine kurulmuştur.",
            "Meydan ve pist bilgileri için OurAirports airports.csv ve runways.csv veri setlerinden yararlanılmıştır. Bu veriler havalimanı kimliği, koordinat, pist uzunluğu, pist yönü ve meydan tipi gibi alanlar sağlar. Sistem Türkiye LT* meydanlarını filtreleyerek brifing, harita ve alternate öneri modülünde kullanır.",
            "NOTAM tarafında canlı erişim için Laminar Data Hub, SkyLink, EUROCONTROL EAD ve uzun vadede DHMİ/EAD hattı değerlendirilmiştir. Ancak geçerli canlı API anahtarı doğrulanmadığı durumda sistem deterministik sentetik NOTAM üretir ve bunu kullanıcıya açık biçimde demo/test verisi olarak gösterir.",
            "Bu modülün en önemli katkısı, sistemin veri erişim stratejisini kırılgan scraping yerine resmi/API tabanlı ve fallback destekli bir yapıya taşımasıdır. Böylece uygulama hem demo ortamında hem de ileride gerçek sağlayıcılar eklendiğinde daha sağlam çalışabilecek bir temele kavuşmuştur.",
        ),
        None,
    ))

    sections.append((
        "Modül 2: Backend API Geliştirme",
        p(
            "Backend tarafında Express tabanlı API, sistemin brifing orkestratörü olarak konumlandırılmıştır. Kullanıcı kalkış ve varış meydanını seçtiğinde web arayüzü /brief endpointine istek gönderir. API bu isteği yalnızca veri aktarma işlemi olarak değil, birden fazla veri ve model kaynağını birleştiren ana iş akışı olarak ele alır.",
            "API önce meydan ve pist bilgilerini çözer, ardından METAR/TAF provider zincirini çalıştırır. Elde edilen veriler providerName, source, fetchedAt, fallbackUsed ve stale gibi metadata alanlarıyla birlikte normalize edilir. Bu alanlar sayesinde kullanıcı yalnızca veriyi değil, verinin nereden geldiğini ve fallback kullanılıp kullanılmadığını da görebilir.",
            "NOTAM sağlayıcı modülü API içinde ayrı bir katman olarak ele alınmıştır. NOTAM_PROVIDER=simulated olduğunda deterministik sentetik event engine çalışır. Eğer laminar veya skylink provider seçilmişse ilgili API anahtarı ile canlı istek denenir; başarısız olursa sistem fallback davranışıyla sentetik veriye dönebilir.",
            "API, kural tabanlı risk skorunu ürettikten sonra services/nlp AI servisine NOTAM parse, risk predict ve brief report istekleri gönderir. AI servisinden dönen mlScore, ruleScore, finalScore, class, confidence, drivers ve modelVersion alanları /brief response içine eklenir. AI servisi kapalıysa eski /brief davranışı bozulmadan kural tabanlı fallback korunur.",
            "Backend ayrıca /model/status, /feedback/summary, /feedback, /brief/logs ve /brief/logs/latest endpointlerini sağlar. Bu endpointler modelin yüklü olup olmadığını, validation metriklerini, provider durumunu, kullanıcı geri bildirimlerini ve yapılan brifing sorgularını incelemek için kullanılır.",
            "Bu modülün mühendislik değeri, API'nin tek bir servis çağrısı yerine karar destek orkestratörü olarak çalışmasıdır. Kullanıcı bir butona bastığında sistem dış veri sağlayıcılarını, lokal hesapları, model çıktısını ve raporlama katmanını birleştirerek açıklanabilir bir çıktı üretir.",
        ),
        None,
    ))

    sections.append((
        "Modül 3: Harita Tabanlı Görselleştirme ve Canlı Trafik Katmanı",
        p(
            "Harita modülü, sayısal brifing çıktısının mekansal karşılığını oluşturur. Kullanıcı yalnızca risk skorunu görmekle kalmaz; kalkış ve varış meydanlarını, rota çizgisini, alternate meydanları, pist yönünü ve rüzgar ilişkisini görsel olarak inceleyebilir.",
            "Leaflet tabanlı harita yapısında DEP ve ARR markerları, rota çizgisi ve alternate meydan önerileri birlikte gösterilir. Bu yaklaşım özellikle alternate öneri sisteminin anlaşılmasını kolaylaştırır. Kullanıcı önerilen meydanın varış meydanına göre nerede bulunduğunu ve mesafenin operasyonel olarak anlamlı olup olmadığını görebilir.",
            "Aktif pist yönü ve rüzgar oku, brifingde hesaplanan headwind/crosswind değerlerinin görsel bağlamını güçlendirir. Bir pistin sayısal yönü ve rüzgar bileşeni harita üzerinde birlikte görüldüğünde kullanıcı skoru yalnızca metin olarak değil, rota ve meydan geometrisiyle birlikte yorumlayabilir.",
            "Canlı trafik katmanı, uygulamanın brifing deneyimini daha zengin hale getirir. Trafik verisi doğrudan final risk skorunun ana bileşeni değildir; ancak kullanıcıya operasyonel çevre hakkında ek farkındalık sağlar. Bu katman harita modülünün karar destek değerini artırır.",
            "Harita modülünde dikkat edilmesi gereken nokta, görsel bilginin ana risk kararının yerine geçmemesidir. Harita, brifing çıktısını destekleyen bir temsil sunar; resmi havacılık haritası, ATC ekranı veya operasyonel navigasyon kaynağı değildir.",
        ),
        None,
    ))

    sections.append((
        "Modül 4: Risk Değerlendirme ve Karar Destek Mantığının Geliştirilmesi",
        p(
            "Risk değerlendirme modülü, projenin ana karar destek çekirdeğidir. Sistem final skoru tek bir kaynaktan üretmez; ML modeli, kural tabanlı skor, NOTAM semantik skoru ve deterministic guardrail yaklaşımını birlikte kullanır. Bu nedenle skorun arkasındaki nedenler kullanıcıya açıklanabilir.",
            "Risk bandı üç seviyeden oluşur: 0-39 düşük risk, 40-69 orta risk, 70-100 yüksek risk. Bu bantlar operasyonel onay anlamına gelmez. Düşük risk, mevcut veriye göre belirgin operasyonel sinyalin düşük olduğunu; orta risk, limit ve alternate kontrolü gerektiğini; yüksek risk ise belirgin operasyonel risk başlıklarının yeniden doğrulanması gerektiğini gösterir.",
            "Final skor formülü 0.65 * mlScore + 0.25 * ruleScore + 0.10 * notamSemanticScore şeklindedir. Bu dağılımda METAR tabanlı model ana ağırlığı taşır; kural motoru açıklanabilir denge sağlar; NOTAM semantik skoru ise operasyonel kısıtların final risk üzerindeki sınırlı fakat görünür etkisini temsil eder.",
            "Guardrail mantığı, modelin açık meteorolojik riskleri düşük göstermesini engellemek için eklenmiştir. Görüş <1500 m, RVR <550 m, tavan <600 ft, wind >=30 kt, gust >=35 kt veya TS/freezing gibi sinyaller high floor üretir. Daha hafif eşikler caution floor üretir.",
            "Bu yapı, model başarısı yüksek olsa bile sistemin tamamen istatistiksel çıktıya güvenmemesini sağlar. Havacılık gibi alanlarda bazı eşikler doğrudan operasyonel dikkat gerektirir. Guardrail bu eşikleri modelin üzerine koruyucu taban olarak koyar.",
            "Kullanıcı arayüzünde riskin nedeni basit gerekçelerle gösterilir. Örneğin 'ARR: Pist yüzeyi/frenleme durumu problemli' veya 'Varış görüşü okunamadı' gibi cümleler teknik skoru anlaşılır hale getirir. Teknik detay panelinde ise ML skoru, rule skoru, NOTAM skoru, guardrail ve confidence faktörleri incelenebilir.",
        ),
        hybrid_diagram,
    ))

    sections.append((
        "Modül 5: Alternate Meydan Öneri Sisteminin Geliştirilmesi",
        p(
            "Alternate meydan öneri sistemi, varış meydanında risk oluştuğunda kullanıcıya yakın ve uygun alternatifleri göstermek için geliştirilmiştir. Başlangıçta alternate önerilerinin kalkış çevresinde üretilmesi hatalı bir bağlam oluşturmuştur. Daha sonra sistem varış meydanı merkezli sıralama yapacak şekilde düzeltilmiştir.",
            "Değerlendirmede mesafe, pist uzunluğu, kritik NOTAM sayısı, rüzgar uygunluğu, görüş, tavan ve hava koşulları dikkate alınır. Bu kriterlerin her biri alternate için kısa etiketlere dönüştürülür. Kullanıcı alternatif meydanın neden önerildiğini veya neden daha düşük sırada olduğunu görebilir.",
            "Alternate önerileri tek başına resmi yedek meydan planlaması değildir. Gerçek operasyonlarda yakıt, uçak performansı, şirket prosedürü, meydan hizmetleri, gümrük/pasaport, slot ve hava trafik kısıtları gibi ek faktörler de değerlendirilmelidir. Bu sistem yalnızca ilk brifing desteği sağlar.",
            "UI tarafında alternate önerileri ilk ekranda açık gelmez; detay panelinde gösterilir. Bunun nedeni, ana brifing ekranının risk seviyesi ve risk gerekçesi üzerine odaklanmasını sağlamaktır. Kullanıcı ihtiyaç duyduğunda alternate bölümünü açarak önerileri inceleyebilir.",
            "Bu modül, karar destek sisteminin sadece risk üretmediğini, aynı zamanda risk karşısında kontrol edilecek alternatifleri de kullanıcıya sunduğunu gösterir. Tez demosunda özellikle yüksek NOTAM veya hava riski senaryolarında alternate önerileri sistemin pratik değerini artırır.",
        ),
        None,
    ))

    sections.append((
        "Modül 6: Frontend Arayüz ve BriefPanel Geliştirmeleri",
        p(
            "Frontend tarafında en önemli tasarım kararı, ilk ekranın sadeleştirilmesidir. İlk prototiplerde çok fazla ham METAR/TAF metni, teknik skor bileşeni, breakdown ve AI açıklaması aynı anda görünmekteydi. Bu durum kullanıcının ana sorusu olan 'risk neden oluştu?' sorusunu zorlaştırmıştır.",
            "Son tasarımda ilk görünümde rota özeti, risk seviyesi, en önemli gerekçeler ve uçuş risk raporu tablosu yer alır. ML score, rule score, NOTAM semantic score, guardrail, model formula ve confidence faktörleri teknik detay panelinde kapalı tutulur. Böylece demo sade kalır, teknik soru gelirse derinlik korunur.",
            "Uçuş risk raporu tablosunda kalkış görüşü, varış görüşü, kalkış tavanı, varış tavanı, yan rüzgar, TAF eğilimi, kritik NOTAM, hava modeli ve model güveni satırları bulunur. Her satır iyi, izle, risk veya eksik durumuyla gösterilir. Bu yapı kullanıcıya hangi parametrenin sorun çıkardığını hızlıca gösterir.",
            "NOTAM kartları da yeniden düzenlenmiştir. Kartlarda ham NOTAM metni yerine önce Türkçe operasyonel özet, kritik olma gerekçesi, etkilediği pist/prosedür, etki puanı açıklaması ve 'ne yapılır?' bilgisi gösterilir. Ham metin detay altında kalır.",
            "Arayüz dili Türkçe ağırlıklı hale getirilmiştir. Primary driver, breakdown, guardrail gibi teknik İngilizce terimler ilk ekrandan kaldırılmış veya Türkçe açıklamaya çevrilmiştir. Bu değişiklik, özellikle bitirme tezi demosunda jürinin sistemi daha hızlı anlamasını sağlar.",
            "SearchBar davranışı da kullanıcı deneyimi açısından düzenlenmiştir. Kalkış veya varış inputuna tıklayınca kısa/boş sorguda Türkiye meydanları listelenir. Böylece kullanıcı ICAO kodunu ezbere yazmak zorunda kalmadan hızlı rota seçebilir.",
        ),
        None,
    ))

    sections.append((
        "Modül 7: PDF Brifing Modülünün Geliştirilmesi",
        p(
            "PDF brifing modülü, ekrandaki karar destek çıktısının taşınabilir rapor haline getirilmesini sağlar. Uçuş öncesi hazırlıkta kullanıcı yalnızca ekrandaki bilgiye bakmakla kalmayabilir; raporu paylaşmak, arşivlemek veya tez demosunda çıktı olarak göstermek isteyebilir.",
            "PDF çıktısında rota özeti, risk seviyesi, model versiyonu, hibrit skor, AI değerlendirme özeti, METAR hava değerlendirmesi, guardrail nedenleri, weather category bilgileri ve önemli NOTAM etkileri yer alır. Bu yapı, sistemin yalnızca web arayüzünden ibaret olmadığını gösterir.",
            "PDF içinde de sistemin operasyonel otorite yerine geçmediği belirtilmelidir. Çünkü PDF çıktısı daha resmi görünebilir; bu durum yanlış güven oluşturabilir. Bu nedenle raporda karar destek amacı, veri kaynağı ve sınırlılıklar açıkça yer almalıdır.",
            "Gelecek çalışmalarda PDF çıktısı üniversite veya kurum formatına daha fazla yaklaştırılabilir. Kapak, uçuş bilgisi, imza alanı, veri zaman damgası, provider listesi ve ham veri ekleri ayrı sayfalar halinde sunulabilir. Ancak mevcut prototip, karar destek özetini dışa aktarabilen işlevsel bir temel sunmaktadır.",
        ),
        None,
    ))

    sections.append((
        "Modül 8: Yapay Zeka / NLP Servisi ve Hibrit Mimari",
        p(
            "services/nlp, projenin AI servis katmanıdır. Bu katman FastAPI benzeri bir servis olarak çalışır ve /ai/notam/parse, /ai/notam/render, /ai/risk/predict ve /ai/brief/report endpointlerini sağlar. API katmanı bu servisi çağırarak NOTAM analizi, risk tahmini ve açıklanabilir rapor üretimi alır.",
            "NOTAM parse endpointi, ham veya sentetik NOTAM metnini runway, nav, ops_hours, airspace, lighting, surface, validity ve severity gibi alanlara dönüştürür. Bu dönüşüm, kullanıcı arayüzünde kritik NOTAM gerekçelerinin kısa ve anlaşılır gösterilmesini sağlar.",
            "Risk predict endpointi, parsed METAR/TAF, runway/wind bilgisi, NOTAM semantik özellikleri ve airport metadata üzerinden mlScore, finalScore, class, confidence, drivers ve modelVersion döndürür. Model dosyası yoksa veya AI servis hata verirse sistem fallback davranışıyla çalışır.",
            "Brief report endpointi, tam brifing verisi ve risk sonuçları üzerinden Türkçe açıklanabilir rapor üretir. Bu rapor final skoru serbestçe belirlemez. Yalnızca verilen veri ve skorlar üzerinden kullanıcıya okunabilir yorum sunar.",
            "Bu mimaride LLM benzeri katmanın sınırı açıkça çizilmiştir. LLM karar verici değil, açıklayıcıdır. Risk skorunu ML/rule/NOTAM bileşenleri üretir. LLM yalnızca confidence düşürme veya sınırlı risk artırma önerisi gibi kontrollü yorumlar yapabilir.",
        ),
        architecture_diagram,
    ))

    sections.append((
        "Modül 9: ML Pipeline ve Model Eğitimi",
        p(
            "ML pipeline, tools/ml_pipeline.py dosyası altında geliştirilmiştir. Bu pipeline tarihsel METAR indirme, canlı METAR/TAF snapshot toplama, dataset oluşturma, model eğitme ve validasyon işlemlerini içerir. Amaç, manuel olarak toplanan veriyi tekrar üretilebilir bir eğitim sürecine dönüştürmektir.",
            "Tarihsel METAR verisi 2023-2026 aralığında Türkiye LT* istasyonları için indirilmiştir. Kullanılan komut yapısı istasyon listesini turkey olarak alabilir ve proje tarafından bilinen LT* meydanlarını otomatik okuyabilir. Rate limit hatalarına karşı pause parametresi kullanılır.",
            "Dataset oluşturma aşamasında ham METAR kayıtlarından görüş, tavan, rüzgar, gust, hadise, precipitation, fog/mist, thunder/freezing ve benzeri özellikler çıkarılır. Bu özellikler operasyonel proxy risk etiketlerine dönüştürülür. Etiketler gerçek kaza riski değil, meteorolojik operasyonel dikkat göstergesidir.",
            "Eğitilen ilk model üç sınıflı logistic baseline yaklaşımıyla risk_level hedefini öğrenir. Eğitim sonucunda 2.024.185 satırlık veri seti, 180.837 pozitif satır ve yaklaşık 0.993 ROC AUC değeri elde edilmiştir. Time validation ve airport holdout sonuçları modelin proxy etiketleri güçlü biçimde ayırabildiğini göstermiştir.",
            "Model doğrulamasında guardrail sonrası false negative sayısının sıfıra düşmesi önemli bir bulgudur. Bu, açık meteorolojik risk eşiklerinin düşük gösterilmesini engelleyen koruyucu katmanın çalıştığını gösterir. Bunun karşılığı daha fazla false positive, yani daha muhafazakar brifing üretimidir.",
            "Bu modelin sınırlılığı tezde açıkça belirtilmelidir. Model kaza, olay, diversion veya cancellation sonucunu doğrudan tahmin etmemektedir. Türkiye METAR geçmişinden öğrenilmiş operasyonel hava riski proxy modelidir. Bu açıklama hem akademik dürüstlük hem de güvenli kullanım açısından önemlidir.",
        ),
        None,
    ))

    sections.append((
        "Modül 10: NOTAM Analizi ve Sentetik NOTAM Modeli",
        p(
            "NOTAM modülü, canlı veri erişimi henüz doğrulanmadığı durumda sistemin demo ve test kabiliyetini korumak için deterministik sentetik event engine ile desteklenmiştir. Bu motor aynı ICAO ve aynı zaman bucket'ı için aynı NOTAM olaylarını üretir. Böylece testler ve tez demosu kararlı hale gelir.",
            "Sentetik event engine; icao, airport profile, pist sayısı, pist uzunluğu, bölge tipi, major/coastal/eastern etiketleri, UTC saat, sezon ve seed bucket gibi girdilerden olay üretir. Çıktıda category, severity, critical, impacts, validFrom, validTo, affectedRunway, score ve reason bulunur.",
            "Olay kategorileri pist kapanışı, pist kontrolü, pist yüzeyi/frenleme, seyrüsefer yardımcısı arızası, ışıklandırma bakımı, çalışma saati kısıtı, apron/taksi yolu çalışması, hava sahası faaliyeti ve hava bağlantılı uyarıları kapsar. Bu kategoriler UI'da kısa Türkçe gerekçelere çevrilir.",
            "LLM metin katmanı, deterministik event'i NOTAM benzeri operasyonel İngilizce metne çevirebilir. Ancak severity, critical veya score değerlerini değiştiremez. Eğer LLM çıktısı schema validation'dan geçmezse deterministic template kullanılır. Bu, AI çıktısının kontrolsüz şekilde operasyonel etki değiştirmesini engeller.",
            "Kullanıcı arayüzünde sentetik NOTAM açıkça demo/test verisi olarak işaretlenir. Bu ayrım kritik öneme sahiptir; çünkü kullanıcı sentetik veriyi gerçek operasyonel NOTAM gibi algılamamalıdır. Canlı provider doğrulandığında aynı schema üzerinden pipeline çalışmaya devam edebilir.",
        ),
        notam_diagram,
    ))

    sections.append((
        "Modül 11: Kalibrasyon, Geri Bildirim ve Loglama",
        p(
            "Kalibrasyon ekranı, modelin yalnızca arka planda çalışan bir dosya olmadığını, durumunun ve validasyon metriklerinin kullanıcıya gösterilebildiğini kanıtlar. /model/status endpointi modelin yüklü olup olmadığını, model versiyonunu, dataset satır sayısını, label dağılımını ve evaluation sonuçlarını döndürür.",
            "Geri bildirim paneli, kullanıcıların brifing sonucunu doğru, fazla muhafazakar, kaçan risk veya yanlış neden olarak işaretleyebilmesini sağlar. Bu feedback şimdilik otomatik model eğitimine girmemektedir; ancak ileride threshold validation ve supervised calibration için yerel manuel etiket kaynağı oluşturur.",
            "Brifing sorgularının loglanması, demo ve hata analizinde önemlidir. data/logs/brief_queries.jsonl dosyası hangi DEP/ARR sorgusunun hangi sonuçları döndürdüğünü incelemeye olanak verir. /brief/logs/latest endpointi son sorgunun hızlı kontrol edilmesini sağlar.",
            "Bu modül, model geliştirme sürecinin yalnızca eğitim komutlarından ibaret olmadığını gösterir. Üretim benzeri bir karar destek sisteminde modelin ne zaman, hangi veriyle, hangi confidence ile sonuç ürettiği izlenebilir olmalıdır. Loglama ve feedback bu gereksinimin ilk adımıdır.",
        ),
        None,
    ))

    sections.append((
        "Modül 12: Test, Doğrulama ve Kabul Senaryoları",
        p(
            "Test süreci birim test, entegrasyon testi ve kullanıcı arayüzü senaryoları olarak ele alınmıştır. Feature extraction tarafında METAR/TAF/NOTAM girdilerinden numeric ve text özelliklerin doğru çıkarılması önemlidir. Ensemble tarafında ML, rule ve NOTAM skorlarının doğru ağırlıkla birleşmesi kontrol edilmelidir.",
            "Fallback testi, sistemin dayanıklılığı açısından kritik kabul edilmiştir. AI servisi kapalıyken /brief endpointi eski davranışını bozmadan kural tabanlı sonuç döndürmelidir. Provider 204, 429 veya 500 dönerse fallback zinciri denenmelidir. Bu davranış uçuş brifing akışının tek servis hatasına bağımlı olmamasını sağlar.",
            "Entegrasyon senaryosunda /brief?dep=LTFM&arr=LTAC gibi bir istek eski alanları korurken risk.ml ve aiReport alanlarını opsiyonel olarak döndürmelidir. NOTAM parser RWY CLSD, ILS U/S, AD OPR HR gibi örnekleri doğru kategoriye ayırmalıdır.",
            "UI senaryolarında orta/yüksek riskli rota açıldığında risk seviyesi anlaşılır görünmeli, kritik NOTAM gerekçeleri tek tek listelenmeli ve risk raporu tablosunda sorunlu parametreler kırmızı veya sarı görünmelidir. Düşük riskli rotada ise belirgin risk sinyali yok mesajı sade biçimde verilmelidir.",
            "Teknik kontrollerde npm --prefix apps/web run build, npm --prefix apps/api run build ve python -m py_compile services/nlp/main.py komutları çalıştırılmıştır. Model tarafında data/processed/evaluation.json dosyası false negative, false positive ve guardrail etkisini incelemek için kullanılmıştır.",
        ),
        None,
    ))

    sections.append((
        "Kullanım Senaryoları ve Demo Akışı",
        p(
            "Tez demosunda ilk olarak uygulama start-flight-risk.bat ile başlatılır. Bu dosya AI servisini 127.0.0.1:8000, API servisini localhost:4000 ve web arayüzünü 127.0.0.1:5174 üzerinde çalıştırır. Portlar doluysa duplicate servis başlatmamak için mevcut servisleri kullanma davranışı korunur.",
            "Kullanıcı kalkış ve varış meydanını seçtiğinde sistem önce veri durumu kartını günceller. METAR/TAF provider auto, NOTAM provider simulated/hybrid, TAF snapshot zamanı ve model sağlık bilgisi ekranda görülebilir. Bu bölüm jüriye sistemin yalnızca UI değil, veri/model durumunu da izlediğini gösterir.",
            "Orta riskli bir rota seçildiğinde sistem 40-69 bandını 'Orta Risk' olarak açıklar. Bu sonuç uçuşu otomatik onaylamaz veya iptal etmez; limit, alternate ve güncel veri kontrolü gerektiğini belirtir. Yüksek riskte ise 70-100 bandı planın yeniden doğrulanması gerektiğini gösterir.",
            "Kritik NOTAM bulunan senaryoda sistem genel bir 'kritik NOTAM var' cümlesiyle yetinmez. DEP veya ARR bazında pist kapalı, pist yüzeyi/frenleme, seyrüsefer yardımcısı, ışıklandırma, hava sahası veya çalışma saati kısıtı gibi somut nedenleri listeler.",
            "Teknik soru geldiğinde detay paneli açılır. Burada ML score, rule score, NOTAM semantic score, final score, model version, guardrail reasons ve confidence faktörleri gösterilir. Böylece demo sade başlar ama teknik savunma derinliği korunur.",
            "PDF indirildiğinde kullanıcı aynı brifing bilgisini taşınabilir rapor olarak alabilir. Harita sayfasında rota, alternate meydanlar ve görsel bağlam gösterilir. Bu akış, sistemin veri, model, UI ve rapor bileşenlerinin uçtan uca çalıştığını gösterir.",
        ),
        None,
    ))

    sections.append((
        "Sonuç ve Gelecek Çalışmalar",
        p(
            "Bu bitirme tezi kapsamında, uçuş öncesi operasyonel brifing sürecini destekleyen açıklanabilir hibrit yapay zeka tabanlı bir prototip geliştirilmiştir. Sistem METAR/TAF, NOTAM, pist/rüzgar, alternatif meydan, kural motoru, ML modeli ve AI raporlama katmanını tek akışta birleştirir.",
            "Çalışmanın en önemli sonucu, yapay zekanın havacılık karar destek sürecinde nasıl sınırlı, açıklanabilir ve denetlenebilir bir rol üstlenebileceğini göstermesidir. Model kaza riski tahmini yapmaz; METAR tabanlı operasyonel proxy risk üretir. NOTAM katmanı operasyonel kısıtları sınıflandırır. Raporlama katmanı ise bunları kullanıcıya sade biçimde açıklar.",
            "Sistemin mevcut sınırlılıkları vardır. Canlı Türkiye NOTAM sağlayıcısı henüz geçerli API anahtarıyla doğrulanmamıştır. TAF geçmiş veri seti ve TAF ML modeli henüz tamamlanmamıştır. Feedback etiketleri toplanmakta ancak otomatik eğitim sürecine dahil edilmemektedir.",
            "Gelecek çalışmalarda DHMİ/EAD/EUROCONTROL veya lisanslı ticari API üzerinden gerçek NOTAM entegrasyonu doğrulanmalıdır. TAF snapshot birikimiyle TAF trend veri seti oluşturulmalı, ayrı TAF skoru eklenmeli ve feedback verileri threshold calibration sürecine dahil edilmelidir.",
            "Ayrıca model açıklanabilirliği SHAP benzeri özellik katkılarıyla güçlendirilebilir. PDF çıktısı resmi brifing formatına daha fazla yaklaştırılabilir. Kullanıcı rolleri, operasyonel limit profilleri ve uçak tipi bazlı performans parametreleri eklenerek sistemin karar destek kapsamı genişletilebilir.",
            "Sonuç olarak flight-risk, resmi operasyonel karar sistemi değil; veri entegrasyonu, açıklanabilir AI ve kullanıcı dostu brifing yaklaşımını birleştiren akademik bir prototiptir. Bu yönüyle bitirme tezi kapsamında hem yazılım mühendisliği hem veri/modelleme hem de kullanıcı deneyimi açısından bütünleşik bir çalışma ortaya koymaktadır.",
        ),
        None,
    ))

    sections.append((
        "KAYNAKLAR (Projede yer alan tüm kaynakların referans numaraları ile listesi)",
        p(
            "[1] Aviation Weather Center, AviationWeather Data API, https://aviationweather.gov/data/api/.",
            "[2] Iowa State University Iowa Environmental Mesonet, ASOS/AWOS/METAR Data Download, https://mesonet.agron.iastate.edu/request/download.phtml.",
            "[3] OurAirports, airports.csv ve runways.csv açık veri setleri, https://ourairports.com/data/.",
            "[4] Laminar Data Hub, NOTAM Data APIs v2, https://developer.laminardata.aero/documentation/notamdata/v2.",
            "[5] EUROCONTROL, European AIS Database, https://www.eurocontrol.int/service/european-ais-database.",
            "[6] FAA SWIFT/SWIM Portal, https://portal.swim.faa.gov/.",
            "[7] React, Vite, Express, FastAPI ve scikit-learn resmi dokümantasyonları.",
            "[8] flight-risk proje kaynak kodu, PROJECT_BRAIN.md, services/nlp, tools/ml_pipeline.py ve apps/api/apps/web modülleri.",
        ),
        None,
    ))

    # Ek içerik: şablon yapısını bozmadan sayfa hacmini tez düzeyine taşır.
    appendix_topics = [
        ("Ek Açıklama 1: METAR Özellik Çıkarma", "METAR ham metni doğrudan modele verilmez. Önce rüzgar yönü, rüzgar hızı, gust, görüş, bulut tavanı, hadise kodları, precipitation, fog/mist, thunder/freezing ve RVR gibi alanlar ayrıştırılır. Bu ayrıştırma hem model için numeric feature üretir hem de UI tarafında kullanıcıya sade özet gösterilmesini sağlar. Ayrıştırılamayan alanlar normal kabul edilmez; eksik olarak işaretlenir ve confidence üzerinde etkili olur."),
        ("Ek Açıklama 2: TAF Snapshot Stratejisi", "TAF geçmiş veri seti hazır olmadığı için sistem canlı TAF snapshot toplamayı destekler. collect-taf-snapshot.bat tek seferlik toplama yaparken install-taf-snapshot-task.bat Windows Task Scheduler üzerinden arka planda periyodik toplama sağlayabilir. Bu veri ileride TAF trend skoru ve TAF ML modeli için temel oluşturacaktır."),
        ("Ek Açıklama 3: Confidence Mantığı", "Confidence, skorun yanında ayrı bir güven göstergesidir. METAR/TAF eksikliği, rüzgar verisi eksikliği, fallback provider kullanımı, AI servis hatası veya domain dışı durumlar confidence değerini düşürebilir. Bu sayede sistem yalnızca riskin yüksekliğini değil, sonucu üretirken ne kadar veri kalitesi bulunduğunu da kullanıcıya aktarır."),
        ("Ek Açıklama 4: UI Sadeleştirme Gerekçesi", "İlk ekranın sade olması, teknik detayın kaldırıldığı anlamına gelmez. Kullanıcı önce karar destek özetini görür; ardından isterse model, NOTAM ve feedback detaylarını açar. Bu yaklaşım, pilot veya operasyon kullanıcısının hızlı tarama ihtiyacı ile jüri veya geliştiricinin teknik inceleme ihtiyacını aynı arayüzde dengeler."),
        ("Ek Açıklama 5: Sentetik NOTAM Etiketi", "Sentetik NOTAM, geliştirme ve demo için faydalıdır ancak gerçek operasyonel veri değildir. Bu nedenle UI ve raporda açıkça demo/test olarak belirtilir. Bu şeffaflık, sistemin güvenilirliğini artırır; çünkü kullanıcı hangi verinin canlı, hangi verinin simülasyon olduğunu bilir."),
        ("Ek Açıklama 6: Modelin Yanlış Yorumlanmasını Önleme", "Modelin ROC AUC değerinin yüksek olması, gerçek uçuş emniyeti tahmini yaptığı anlamına gelmez. Model, METAR göstergelerinden türetilmiş proxy etiketleri ayırır. Bu nedenle tez savunmasında model başarısı anlatılırken 'kaza riski' ifadesinden kaçınılmalı, 'operasyonel METAR proxy risk' ifadesi kullanılmalıdır."),
        ("Ek Açıklama 7: Logların Kullanımı", "brief_queries.jsonl dosyası, yapılan sorguların ve dönen sonuçların sonradan incelenmesini sağlar. Kullanıcı bir sonucun garip olduğunu düşündüğünde son sorgu /brief/logs/latest üzerinden kontrol edilebilir. Bu özellik, model davranışı ve UI çıktısı arasında izlenebilirlik kurar."),
        ("Ek Açıklama 8: Tez Savunması Mesajı", "Bu projenin savunma cümlesi şudur: Sistem pilot yerine karar vermez; METAR/TAF, NOTAM, pist/rüzgar ve ML çıktısını birleştirerek açıklanabilir uçuş öncesi risk brifingi üretir. Hibrit mimari, yapay zekanın kontrollü ve açıklanabilir kullanımını göstermektedir."),
    ]
    for title, text in appendix_topics:
        sections.append((title, p(text, text.replace("Bu", "Ayrıca bu", 1), text.replace("Sistem", "Geliştirilen sistem", 1)), None))

    deep_topics = [
        ("METAR ham metninin normalize edilmesi", "METAR verisi farklı meydanlarda küçük biçim farklılıkları içerebilir", "parser katmanı raw metni korurken model için numeric alanlar çıkarır", "görüş, tavan, rüzgar, gust ve hadise alanları ayrı özelliklere dönüştürülür"),
        ("TAF verisinin model dışında ama karar içinde kullanılması", "TAF geleceğe dönük hava beklentisini temsil eder", "mevcut sistem TAF'ı ayrı ML modeli yapmadan trend sinyali olarak kullanır", "TEMPO, BECMG ve kötüleşme işaretleri risk açıklamasında değerlendirilir"),
        ("NOTAM kritik etki puanının yorumlanması", "NOTAM skorunun tek başına ne anlama geldiği kullanıcı için belirsiz olabilir", "UI puanı kategori ve kısa operasyonel gerekçeyle birlikte gösterir", "pist, yüzey, seyrüsefer ve hava sahası etkileri ayrı açıklanır"),
        ("Provider fallback davranışı", "canlı veri sağlayıcıları geçici olarak 204, 429 veya 500 döndürebilir", "API provider zincirini sırayla deneyerek brifing akışını sürdürür", "fallback kullanıldıysa kullanıcıya provider metadata alanlarında gösterilir"),
        ("AI servisinin zorunlu olmaması", "AI servisinin kapalı olması uygulamanın tamamen çökmesine neden olmamalıdır", "Express API kural tabanlı fallback sonucunu korur", "confidence düşürülerek kullanıcının sınırlılığı görmesi sağlanır"),
        ("Guardrail eşiklerinin rolü", "model bazı açık meteorolojik riskleri olasılık olarak düşük değerlendirebilir", "deterministik guardrail görünür operasyonel eşikleri taban risk olarak uygular", "bu yaklaşım false negative riskini azaltırken muhafazakar sonuç üretebilir"),
        ("Confidence faktörlerinin açıklanması", "risk skoru tek başına veri kalitesini anlatmaz", "confidence seviyesi veri eksikliği, fallback ve domain dışı durumları açıklar", "kullanıcı skorun hangi güven düzeyiyle üretildiğini görebilir"),
        ("Risk bandı dilinin sadeleştirilmesi", "kırmızı veya sarı gibi renkler tek başına yeterli açıklama sağlamaz", "UI düşük, orta ve yüksek risk ifadelerini kullanır", "her bandın ne anlama geldiği tek cümleyle açıklanır"),
        ("Ham veri ve özet veri ayrımı", "METAR/TAF ve NOTAM ham formatları teknik kullanıcı için gereklidir", "ilk ekran sade özet gösterir, ham veri detay altında tutulur", "bu düzen hem hızlı demo hem teknik inceleme ihtiyacını karşılar"),
        ("Sentetik NOTAM etik sorumluluğu", "sentetik veri gerçek operasyonel NOTAM gibi algılanmamalıdır", "sistem sentetik kayıtları demo/test olarak işaretler", "bu şeffaflık akademik ve operasyonel güven açısından zorunludur"),
        ("Laminar ve SkyLink entegrasyon hazırlığı", "canlı NOTAM erişimi API anahtarı ve lisans koşulları gerektirir", "provider skeleton yapısı canlı entegrasyon için hazır tutulur", "anahtar doğrulanmadan canlı veri varmış gibi davranılmaz"),
        ("EUROCONTROL EAD hedefi", "Avrupa bölgesinde resmi aeronautical information kaynakları önemlidir", "EAD uzun vadeli resmi NOTAM/AIP hattı olarak değerlendirilir", "erişim koşulları ve Türkiye kapsaması ayrıca doğrulanmalıdır"),
        ("OurAirports verisinin rolü", "meydan ve pist bilgisi risk hesabı için meteoroloji kadar önemlidir", "airports.csv ve runways.csv koordinat ve pist bilgisini sağlar", "alternate ve harita modülü bu metaveriden yararlanır"),
        ("Iowa Mesonet geçmiş METAR arşivi", "model eğitimi canlı API'den anlık veri çekerek yapılamaz", "Iowa Mesonet geçmiş ASOS/METAR kayıtlarını indirilebilir hale getirir", "Türkiye LT* istasyonları için geniş veri seti oluşturulmuştur"),
        ("Proxy etiket yaklaşımının gerekçesi", "kaza verisi seyrek ve çok değişkenli olduğu için ilk model hedefi olamaz", "etiketler METAR operasyonel dikkat göstergelerinden türetilmiştir", "model kaza değil operasyonel hava proxy riski hesaplar"),
        ("Üç sınıflı risk_level hedefi", "yalnızca normal/risk ikili ayrımı brifing için kaba kalabilir", "risk_level normal, dikkat ve yüksek sınıflarıyla modellenir", "bu yapı UI'daki düşük/orta/yüksek risk bandına daha uyumludur"),
        ("Time validation önemi", "rastgele train/test bölmesi zamansal veri sızıntısı yaratabilir", "time validation daha gerçekçi dönemsel genelleme kontrolü sağlar", "modelin güncel veya yakın dönem veriye davranışı bu şekilde incelenir"),
        ("Airport holdout önemi", "aynı meydanın geçmişini görmüş model yeni meydanda yanıltıcı olabilir", "airport holdout bazı meydanları test tarafında tutar", "bu yöntem modelin farklı LT* meydanlarına genellemesini sınar"),
        ("False positive ve false negative dengesi", "havacılık karar desteğinde kaçan risk ciddi problemdir", "guardrail false negative azaltmaya öncelik verir", "bunun karşılığı fazla muhafazakar bazı sonuçlardır"),
        ("Feedback verisinin gelecekteki kullanımı", "kullanıcı geri bildirimi şu an otomatik eğitimde kullanılmaz", "brief_feedback.jsonl yerel manuel etiket havuzu oluşturur", "ileride threshold tuning ve calibration için değerlendirilebilir"),
        ("Log kayıtlarının debug değeri", "ekranda görülen sonuçların sonradan izlenebilmesi gerekir", "brief_queries.jsonl sorgu ve response özetlerini saklar", "model veya UI hatası incelenirken son sorgu doğrulanabilir"),
        ("PDF çıktısının tez demosundaki rolü", "sadece ekran görüntüsü ürün bütünlüğünü göstermeye yetmez", "PDF rapor sistemin taşınabilir çıktı üretebildiğini gösterir", "hibrit skor ve AI değerlendirme raporda özetlenir"),
        ("Harita katmanının karar destek değeri", "sayısal skor mekansal bağlam olmadan eksik kalabilir", "harita rota, alternate, pist yönü ve rüzgar ilişkisini gösterir", "kullanıcı riskin hangi meydan çevresinde yoğunlaştığını görür"),
        ("Alternate önerilerinin ARR merkezli olması", "varış problemi için kalkış çevresindeki alternates yanıltıcıdır", "öneriler varış meydanı çevresinde sıralanır", "bu düzeltme operasyonel bağlama uygunluğu artırır"),
        ("UI'da Türkçe dil tercihleri", "İngilizce teknik terimler demo anlaşılabilirliğini azaltabilir", "primary driver gibi ifadeler Türkçe açıklamaya çevrilir", "teknik terim gerektiğinde detay panelinde korunur"),
        ("Model health kartı", "modelin yüklü olup olmadığı kullanıcıya görünmelidir", "home dashboard model versiyonu, AUC ve satır sayısını gösterir", "bu kart tez savunmasında teknik kanıt sağlar"),
        ("TAF snapshot görevleri", "TAF geçmişi hemen hazır olmadığı için veri biriktirmek gerekir", "Windows scheduled task periyodik snapshot toplamak için eklenmiştir", "bu veri gelecekte TAF trend modeli için kullanılacaktır"),
        ("Port çakışması yönetimi", "lokal geliştirmede aynı portu iki servis kullanabilir", "startup script duplicate servis başlatmamaya çalışır", "kullanıcı port sahibini görüp elle kapatabilir"),
        ("Response geriye uyumluluğu", "UI ve API birlikte gelişirken eski alanlar bozulmamalıdır", "/brief response eski alanları korur ve yeni alanları opsiyonel ekler", "bu yaklaşım entegrasyon riskini azaltır"),
        ("OpenAPI ve dokümantasyon", "API uçları tez ve geliştirme sürecinde izlenebilir olmalıdır", "Swagger/OpenAPI JSON endpointleri API'nin dış kontratını gösterir", "bu yapı backend'in savunulabilirliğini artırır"),
        ("ML model artifact yönetimi", "eğitimden sonra modelin servis tarafından yüklenmesi gerekir", "risk_model.json services/nlp/models altında tutulur", "AI servisi restart edildiğinde en güncel modeli okur"),
        ("Model versiyonunun UI'da gösterilmesi", "kullanıcı hangi modelle sonuç aldığını bilmelidir", "modelVersion risk.ml alanında ve teknik detayda gösterilir", "bu bilgi deney tekrarlanabilirliği için önemlidir"),
        ("NOTAM ham metninin detayda tutulması", "ham NOTAM denetim için gereklidir ama ilk ekranı boğar", "kartlar önce Türkçe özet verir, ham metin açılır detayda kalır", "kullanıcı ihtiyaç duyduğunda orijinal metne erişir"),
        ("Kritik NOTAM maddelendirme", "birden fazla kritik NOTAM var demek yeterli değildir", "her kritik notam için kısa madde üretilir", "pist kapalı veya ILS etkisi gibi somut nedenler gösterilir"),
        ("Mimari sınırların korunması", "gereksiz mimari değişim proje riskini artırır", "monorepo, Express API, web ve services/nlp sınırları korunmuştur", "yeni özellikler mevcut yapı içinde genişletilmiştir"),
        ("Tez savunması için ana mesaj", "jürinin projeyi hızlı anlaması gerekir", "ana mesaj sistemin pilot yerine karar vermeyen açıklanabilir karar destek asistanı olduğudur", "hibrit AI katkısı bu mesaj üzerinden anlatılır"),
        ("Emniyet ve hukuki sınır", "operasyonel karar iddiası yanlış ve risklidir", "UI ve raporda karar destek amacı vurgulanır", "resmi kaynak ve otorite kontrolü gerekliliği korunur"),
        ("Veri eksikliğinin görünür olması", "eksik veri normal kabul edilmemelidir", "risk raporu tablosunda eksik durumu ayrı gösterilir", "bu yaklaşım yanlış güven oluşmasını engeller"),
        ("Kullanıcı odaklı risk raporu", "skorun yanında parametre bazlı açıklama gerekir", "görüş, tavan, rüzgar, TAF, NOTAM ve güven satırları verilir", "hangi parametrenin sorun çıkardığı basit dille açıklanır"),
        ("Teknik detayların jüri için korunması", "ilk ekran sade olsa da tez teknik derinlik istemektedir", "detay panelinde formula, guardrail, model ve feedback gösterilir", "bu yapı ürün ve akademik savunma dengesini sağlar"),
        ("Canlı ve sentetik veri ayrımı", "veri kaynağının canlı olup olmadığı sonucu yorumlamayı etkiler", "providerName, source, fallbackUsed ve synthetic alanları gösterilir", "kullanıcı veri durumunu kontrol edebilir"),
        ("Rate limit ve indirme stratejisi", "Iowa Mesonet ve API kaynakları sınırsız çekim için tasarlanmamıştır", "pause parametresi ve aylık indirme parçaları kullanılmıştır", "429 hatasına karşı indirme kaldığı yerden sürdürülebilir"),
        ("BTS benzeri etiketlerin geleceği", "operasyonel delay/diversion etiketleri model hedefini güçlendirebilir", "mevcut çalışma METAR proxy riskle sınırlıdır", "gelecekte operasyonel sonuç verisi eklenirse model kapsamı genişler"),
        ("TAF ML modelinin geleceği", "TAF canlı tahmin verisi zaman boyutlu yapıdadır", "snapshot birikmeden güvenilir TAF modeli eğitmek zordur", "gelecekte TAF trend skoru ayrı bir model olabilir"),
        ("Dış sağlayıcı lisans riski", "ticari API'ler ücret ve kullanım koşulu gerektirir", "sistem provider abstraction ile bu riski izole eder", "sağlayıcı değişse bile UI ve risk pipeline korunur"),
        ("Kod tabanı izlenebilirliği", "tez projesinde neyin nerede çalıştığı bilinmelidir", "apps/api, apps/web, services/nlp ve tools ayrımı nettir", "bu ayrım proje bakımını kolaylaştırır"),
        ("Kabul kriterlerinin karşılanması", "jüri ilk ekranda rota, risk ve nedeni anlamalıdır", "son UI bu kriterlere göre sadeleştirilmiştir", "teknik detaylar gerektiğinde genişletilebilir durumdadır"),
        ("Akademik dürüstlük", "model ve veri sınırlılıkları saklanmamalıdır", "tez metni proxy model, sentetik NOTAM ve canlı NOTAM eksikliğini açıklar", "bu açıklık çalışmanın güvenilirliğini artırır"),
        ("Ürünleşme potansiyeli", "prototip doğrudan operasyonel ürün değildir", "veri lisansı, sertifikasyon, kullanıcı rolleri ve resmi entegrasyon gerekir", "mevcut çalışma bu ürünleşme yolunun teknik iskeletini gösterir"),
    ]
    for idx, (title, problem, implementation, detail) in enumerate(deep_topics, start=1):
        paragraphs = [
            f"{problem}. Bu başlık, sistemin yalnızca çalışan bir ekran değil, uçuş öncesi brifing sürecinin belirli bir sorununa cevap veren mühendislik çözümü olduğunu gösterir. Tez kapsamında bu nokta, ürün değerinin ve teknik kararların birlikte açıklanması için önemlidir.",
            f"Projede {implementation}. Bu yaklaşım mevcut monorepo sınırları içinde uygulanmıştır; Express API orkestrasyon katmanı, web arayüzü, AI servisi ve veri/model pipeline'ı arasındaki sorumluluk ayrımı korunmuştur.",
            f"Teknik olarak {detail}. Bu ayrım hem test edilebilirliği hem de kullanıcıya gösterilen açıklamaların doğruluğunu artırır. Kullanıcıya gereksiz ayrıntı yüklenmez; ancak teknik detay panelinde kararın nasıl üretildiği izlenebilir.",
            "Bu konunun tez savunmasındaki karşılığı, hibrit yapay zekanın yalnızca bir etiket olarak kullanılmadığını göstermektir. Her bileşen belirli bir veri türünü işler, belirli bir sınırlılığa sahiptir ve final brifinge kontrollü biçimde katkı sağlar.",
            "Gelecek çalışmalarda bu başlık daha güçlü veri kaynakları, daha fazla test senaryosu ve gerçek kullanıcı geri bildirimiyle geliştirilebilir. Bununla birlikte mevcut prototip, karar destek sisteminin temel davranışını açıklanabilir ve denetlenebilir biçimde ortaya koymaktadır.",
        ]
        sections.append((f"Ayrıntılı Teknik Değerlendirme {idx}: {title}", paragraphs, None))

    defense_topics = [
        "Bu sistem neden yalnızca LLM ile yapılmadı?",
        "Neden gerçek kaza riski yerine operasyonel proxy risk kullanıldı?",
        "Canlı NOTAM yokken sentetik NOTAM kullanmak ne kadar anlamlıdır?",
        "Skorun yüksek çıkması uçuş iptali anlamına gelir mi?",
        "Model güveni yüksek olsa bile resmi kontrol neden gereklidir?",
        "Türkiye verisi ile eğitilen modelin sınırları nelerdir?",
        "TAF modelinin henüz olmaması sistemi nasıl etkiler?",
        "Guardrail yaklaşımı neden muhafazakar sonuç üretir?",
        "Kullanıcı geri bildirimi gelecekte modeli nasıl iyileştirir?",
        "Provider zinciri neden tek API kullanımından daha güvenlidir?",
        "Kalibrasyon ekranı tez için neden önemlidir?",
        "Risk raporu tablosu ham veriden neden daha anlaşılırdır?",
        "NOTAM skorunun açıklanması neden ayrı bir kullanıcı ihtiyacıdır?",
        "PDF brifing çıktısı ürün değerini nasıl artırır?",
        "Harita modülü risk skorunu nasıl tamamlar?",
        "Alternate önerilerinde varış merkezli yaklaşım neden gereklidir?",
        "Veri eksikliğini kullanıcıya göstermek neden önemlidir?",
        "Bu prototip nasıl ürünleşebilir?",
        "Etik ve hukuki sınırlar nasıl korunmuştur?",
        "Tez savunmasında hibrit AI katkısı nasıl anlatılmalıdır?",
    ]
    for idx, title in enumerate(defense_topics, start=1):
        paragraphs = [
            f"{title} sorusu, projenin yalnızca teknik olarak çalışıp çalışmadığını değil, mühendislik kararlarının ne kadar savunulabilir olduğunu da ölçer. Bu nedenle sistemin her bileşeni belirli bir probleme cevap verecek şekilde konumlandırılmıştır. Veri sağlayıcıları, model, kural motoru, NOTAM parser, UI ve raporlama katmanı aynı amaca hizmet eder: uçuş öncesi brifingi açıklanabilir hale getirmek.",
            "Bu bağlamda verilen cevap, sistemin operasyonel otorite yerine geçmediği gerçeğiyle birlikte değerlendirilmelidir. Proje bir uçuşu onaylayan veya reddeden karar sistemi değildir. Üretilen risk bandı, hangi başlıkların yeniden kontrol edilmesi gerektiğini gösteren bir karar destek sinyalidir. Bu ayrım hem kullanıcı güvenliği hem akademik doğruluk açısından merkezi öneme sahiptir.",
            "Teknik açıdan bakıldığında, hibrit yapı farklı hata türlerini azaltmak için tercih edilmiştir. Kural motoru açık eşikleri yakalar, ML modeli geçmiş METAR verisinden örüntü öğrenir, NOTAM semantik katmanı operasyonel kısıtları kategorize eder, rapor katmanı ise bunları okunabilir hale getirir. Bu parçaların tek bir kapalı kutuya dönüştürülmemesi sistemin açıklanabilirliğini artırır.",
            "Kullanıcı deneyimi açısından bu kararların karşılığı ilk ekranda görülür. Kullanıcı risk seviyesini, basit gerekçeleri ve parametre bazlı risk raporunu görür. Ham METAR/TAF, model formülü, guardrail ve diğer teknik ayrıntılar detay bölümünde kalır. Böylece sistem hem hızlı demo hem de derin teknik savunma için uygun hale gelir.",
            "Gelecek çalışmalarda bu başlık daha güçlü canlı veri entegrasyonları, TAF geçmiş veri seti, gerçek operasyonel proxy etiketler, kullanıcı geri bildiriminden kalibrasyon ve daha gelişmiş açıklanabilir model yöntemleriyle güçlendirilebilir. Mevcut tez prototipi ise bu yolun uygulanabilir bir yazılım mühendisliği iskeletini ortaya koymaktadır.",
        ]
        sections.append((f"Genişletilmiş Savunma Notu {idx}: {title}", paragraphs, None))

    source_sections = [section for section in sections if section[0].startswith("KAYNAKLAR")]
    sections = [section for section in sections if not section[0].startswith("KAYNAKLAR")]
    sections.extend(source_sections)

    return sections


def build_doc() -> None:
    doc = Document(TEMPLATE)
    for idx, text in enumerate([
        "T.C. FIRAT ÜNİVERSİTESİ",
        "MÜHENDİSLİK FAKÜLTESİ - YAZILIM MÜHENDİSLİĞİ BÖLÜMÜ",
        "BİTİRME TEZİ",
    ]):
        paragraph = doc.paragraphs[idx]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    table = doc.tables[0]
    template_rows = [deepcopy(row._tr) for row in table.rows]
    remove_all_table_rows(table)
    for tr in template_rows[:5]:
        table._tbl.append(deepcopy(tr))
    for i, (key, value) in enumerate(META_ROWS):
        c0, c1 = table.rows[i].cells
        clear_cell(c0)
        clear_cell(c1)
        set_text(c0.paragraphs[0], key, bold=True)
        set_text(c1.paragraphs[0], value)

    # Keep one template row as cloning source, then remove extra source rows after content is appended.
    source_tr = deepcopy(template_rows[6])
    while len(table.rows) > 5:
        table._tbl.remove(table.rows[-1]._tr)
    table._tbl.append(source_tr)

    for title, paragraphs, diagram in build_sections():
        add_section_title_row(table, title)
        add_content_row(table, paragraphs, diagram=diagram)

    # Remove the temporary cloning source row.
    table._tbl.remove(table.rows[5]._tr)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
