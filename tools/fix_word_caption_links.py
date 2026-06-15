from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


TARGET = Path(r"C:/Users/mete_/Downloads/Bitirme_Tez_guncellenmis.docx")
BACKUP = TARGET.with_name("Bitirme_Tez_guncellenmis_before_caption_links.docx")


FIGURE_REFS = [
    ("guardrail akışı", "fig_1", "Şekil", 1, "Hibrit risk skor ve guardrail akışı"),
    ("Veri sağlayıcı zinciri", "fig_2", "Şekil", 2, "Veri sağlayıcı zinciri"),
    ("Brief request sequence", "fig_3", "Şekil", 3, "Brief request sequence"),
    ("AI servis mimarisi", "fig_4", "Şekil", 4, "AI servis mimarisi"),
    ("NOTAM canlı/sentetik pipeline", "fig_5", "Şekil", 5, "NOTAM canlı/sentetik pipeline"),
    ("ML pipeline akışı", "fig_6", "Şekil", 6, "ML pipeline akışı"),
    ("BriefPanel bilgi hiyerarşisi", "fig_7", "Şekil", 7, "BriefPanel bilgi hiyerarşisi"),
    ("Feedback ve kalibrasyon döngüsü", "fig_8", "Şekil", 8, "Feedback ve kalibrasyon döngüsü"),
]

TABLE_REFS_EXISTING = [
    ("Kullanılan teknolojiler", "tbl_1", "Tablo", 1, "Kullanılan teknolojiler"),
    ("Veri kaynakları ve kullanım amacı", "tbl_2", "Tablo", 2, "Veri kaynakları ve kullanım amacı"),
    ("Proxy-risk etiket mantığı", "tbl_3", "Tablo", 3, "Proxy-risk etiket mantığı"),
    ("API endpoint sorumlulukları", "tbl_5", "Tablo", 5, "API endpoint sorumlulukları"),
    ("NOTAM kategori sınıfları", "tbl_6", "Tablo", 6, "NOTAM kategori sınıfları"),
    ("Model eğitim sonuçları", "tbl_7", "Tablo", 7, "Model eğitim sonuçları"),
    ("Validasyon sonuçları", "tbl_8", "Tablo", 8, "Validasyon sonuçları"),
    ("SkyLink canlı NOTAM smoke test sonucu", "tbl_9", "Tablo", 9, "SkyLink canlı NOTAM smoke test sonucu"),
    ("Canlı METAR/TAF ve NOTAM ile uçuş analizi sonuçları", "tbl_10", "Tablo", 10, "Canlı METAR/TAF ve NOTAM ile uçuş analizi sonuçları"),
    ("Sistem kabul kriterleri", "tbl_11", "Tablo", 11, "Sistem kabul kriterleri"),
]

TABLE_REFS = [
    ("tbl_1", "Tablo", 1, "Kullanılan teknolojiler"),
    ("tbl_2", "Tablo", 2, "Veri kaynakları ve kullanım amacı"),
    ("tbl_3", "Tablo", 3, "Proxy-risk etiket mantığı"),
    ("tbl_4", "Tablo", 4, "Risk bandı eşikleri"),
    ("tbl_5", "Tablo", 5, "API endpoint sorumlulukları"),
    ("tbl_6", "Tablo", 6, "NOTAM kategori sınıfları"),
    ("tbl_7", "Tablo", 7, "Model eğitim sonuçları"),
    ("tbl_8", "Tablo", 8, "Validasyon sonuçları"),
    ("tbl_9", "Tablo", 9, "SkyLink canlı NOTAM smoke test sonucu"),
    ("tbl_10", "Tablo", 10, "Canlı METAR/TAF ve NOTAM ile uçuş analizi sonuçları"),
    ("tbl_11", "Tablo", 11, "Sistem kabul kriterleri"),
]


def set_run_font(run, *, bold: bool | None = None, size: float = 12, color: str = "000000") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, align=None, bold_prefix: str | None = None) -> None:
    clear_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        rest = text[len(bold_prefix) :]
        if rest:
            run = paragraph.add_run(rest)
            set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def set_caption(paragraph: Paragraph, label: str, number: int, title: str, *, align) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{label} {number} ")
    set_run_font(run, bold=True)
    run = paragraph.add_run(title)
    set_run_font(run)


def set_list_heading(paragraph: Paragraph, title: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_run_font(run, bold=True, size=14)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "A7A7A7")


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def norm(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).casefold()


def find_paragraph(doc: Document, text: str, *, start: int = 0, exact: bool = False) -> Paragraph:
    needle = norm(text)
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < start:
            continue
        haystack = norm(paragraph.text)
        if (exact and haystack == needle) or (not exact and needle in haystack):
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def find_figure_placeholder(doc: Document, text: str, *, start: int = 0) -> Paragraph:
    needle = norm(text)
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < start:
            continue
        haystack = norm(paragraph.text)
        if needle in haystack and (haystack.startswith("(") or haystack.startswith("şema")):
            return paragraph
    raise ValueError(f"figure placeholder not found: {text}")


def find_table_caption(doc: Document, text: str, *, start: int = 0) -> Paragraph:
    needle = norm(text)
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < start:
            continue
        haystack = norm(paragraph.text)
        if haystack.startswith("tablo") and needle in haystack and haystack.endswith(":"):
            return paragraph
    raise ValueError(f"table caption not found: {text}")


def paragraph_index(doc: Document, target: Paragraph) -> int:
    target_el = target._p
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target_el:
            return idx
    raise ValueError("paragraph index not found")


def list_paragraphs_between(doc: Document, start_heading: Paragraph, end_heading: Paragraph) -> list[Paragraph]:
    start = paragraph_index(doc, start_heading) + 1
    end = paragraph_index(doc, end_heading)
    result = []
    for paragraph in doc.paragraphs[start:end]:
        text = norm(paragraph.text)
        if text and text != "sayfa":
            result.append(paragraph)
    return result


def add_page_header_after(heading: Paragraph) -> None:
    p = insert_paragraph_after(heading)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Sayfa")
    set_run_font(run, bold=True)


def append_hyperlink_run(hyperlink, text: str, *, bold: bool = False) -> None:
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    r.append(r_pr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    hyperlink.append(r)


def add_internal_hyperlink(paragraph: Paragraph, anchor: str, label: str, title: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    append_hyperlink_run(hyperlink, label, bold=True)
    append_hyperlink_run(hyperlink, f" {title}")
    paragraph._p.append(hyperlink)


def add_pageref_field(paragraph: Paragraph, anchor: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {anchor} \\h "
    run._r.append(instr)

    run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    run = paragraph.add_run("1")
    set_run_font(run)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_list_entry(paragraph: Paragraph, anchor: str, label: str, number: int, title: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    add_internal_hyperlink(paragraph, anchor, f"{label} {number}", title)
    run = paragraph.add_run("\t")
    set_run_font(run)
    add_pageref_field(paragraph, anchor)


def existing_bookmark_ids(doc: Document) -> set[int]:
    ids: set[int] = set()
    for node in doc._element.xpath(".//w:bookmarkStart"):
        value = node.get(qn("w:id"))
        if value and value.isdigit():
            ids.add(int(value))
    return ids


def next_bookmark_id(doc: Document) -> int:
    ids = existing_bookmark_ids(doc)
    return max(ids, default=0) + 1


def remove_managed_bookmarks(doc: Document) -> None:
    prefixes = ("fig_", "tbl_")
    ids_to_remove: set[str] = set()
    for node in list(doc._element.xpath(".//w:bookmarkStart")):
        name = node.get(qn("w:name"), "")
        if name.startswith(prefixes):
            ids_to_remove.add(node.get(qn("w:id")))
            node.getparent().remove(node)
    for node in list(doc._element.xpath(".//w:bookmarkEnd")):
        if node.get(qn("w:id")) in ids_to_remove:
            node.getparent().remove(node)


def add_bookmark(paragraph: Paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def insert_table4_caption(doc: Document, bookmark_id: int) -> None:
    # The fourth actual table is the risk-band table and originally has no caption paragraph.
    table = doc.tables[3]
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    paragraph = Paragraph(p, table._parent)
    set_caption(paragraph, "Tablo", 4, "Risk bandı eşikleri", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_bookmark(paragraph, "tbl_4", bookmark_id)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def update_inline_references(doc: Document) -> None:
    replacements = {
        "Tablo 5.3'te": "Tablo 9'da",
        "Tablo 5.3’te": "Tablo 9’da",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not any(old in text for old in replacements):
            continue
        for old, new in replacements.items():
            text = text.replace(old, new)
        set_paragraph_text(paragraph, text, align=paragraph.alignment)


def main() -> None:
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    remove_managed_bookmarks(doc)
    bookmark_id = next_bookmark_id(doc)

    fig_heading = find_paragraph(doc, "ŞEKİL LİSTESİ", exact=True)
    table_heading = find_paragraph(doc, "TABLO LİSTESİ", exact=True)
    abbreviations_heading = find_paragraph(doc, "KISALTMALAR", exact=True)
    fig_list_paras = list_paragraphs_between(doc, fig_heading, table_heading)
    table_list_paras = list_paragraphs_between(doc, table_heading, abbreviations_heading)
    figure_targets = {search: find_figure_placeholder(doc, search, start=100) for search, *_ in FIGURE_REFS}
    table_targets = {search: find_table_caption(doc, search, start=100) for search, *_ in TABLE_REFS_EXISTING}

    set_list_heading(fig_heading, "ŞEKİLLER LİSTESİ")
    set_list_heading(table_heading, "TABLOLAR LİSTESİ")

    add_page_header_after(fig_heading)
    add_page_header_after(table_heading)

    for search, anchor, label, number, title in FIGURE_REFS:
        p = figure_targets[search]
        set_caption(p, label, number, title, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_bookmark(p, anchor, bookmark_id)
        bookmark_id += 1

    for search, anchor, label, number, title in TABLE_REFS_EXISTING:
        p = table_targets[search]
        set_caption(p, label, number, title, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_bookmark(p, anchor, bookmark_id)
        bookmark_id += 1

    insert_table4_caption(doc, bookmark_id)
    bookmark_id += 1

    for p, (anchor, label, number, title) in zip(
        fig_list_paras[: len(FIGURE_REFS)], [(x[1], x[2], x[3], x[4]) for x in FIGURE_REFS]
    ):
        set_list_entry(p, anchor, label, number, title)
    if len(fig_list_paras) > len(FIGURE_REFS):
        remove_paragraph(fig_list_paras[-1])

    for p, (anchor, label, number, title) in zip(table_list_paras, TABLE_REFS):
        set_list_entry(p, anchor, label, number, title)

    update_inline_references(doc)
    set_update_fields_on_open(doc)
    doc.save(TARGET)
    print(f"updated {TARGET}")
    print(f"backup {BACKUP}")


if __name__ == "__main__":
    main()
