from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "Bitirme_Tezi_Kitapcik_Artifact_Genis.md"
OUTPUT = ROOT / "Bitirme_Tezi_Kitapcik_Artifact_Genis.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color="D9DDE3", size="6", space="4") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color in [
        ("Heading 1", 16, "1F2937"),
        ("Heading 2", 14, "26374A"),
        ("Heading 3", 12, "334155"),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str, style: str | None = None, align=None, bold=False) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    elif style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def add_inline_markdown(paragraph, text: str, base_bold=False) -> None:
    # Handles a useful subset: **bold**, `code`, and plain text.
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, bold=base_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=10, color="111827")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, bold=base_bold)


def add_markdown_paragraph(doc: Document, text: str, cover_mode: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cover_mode else WD_ALIGN_PARAGRAPH.JUSTIFY
    if cover_mode:
        p.paragraph_format.space_after = Pt(10)
    add_inline_markdown(p, text)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        raw = lines[i].strip()
        parts = [p.strip() for p in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF7")
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            value = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(p, value, base_bold=(r_idx == 0))
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9, color="111827")
        set_paragraph_border(p, color="EEF2F7", size="2", space="1")
    doc.add_paragraph()


def add_bullet_or_number(doc: Document, line: str, cover_mode=False) -> bool:
    stripped = line.strip()
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, stripped[2:])
        return True
    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
        return True
    return False


def build() -> None:
    doc = Document()
    style_document(doc)

    all_lines = INPUT.read_text(encoding="utf-8").splitlines()
    try:
        start_idx = next(i for i, value in enumerate(all_lines) if value.strip() == "# KAPAK")
        lines = all_lines[start_idx:]
    except StopIteration:
        lines = all_lines
    i = 0
    in_code = False
    code_lines: list[str] = []
    cover_mode = False
    first_h1 = True

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            cover_mode = title == "KAPAK"
            if cover_mode:
                i += 1
                continue
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cover_mode else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(p, title, base_bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[3:].strip(), base_bold=True)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_markdown(p, stripped[4:].strip(), base_bold=True)
            i += 1
            continue

        if add_bullet_or_number(doc, stripped, cover_mode=cover_mode):
            i += 1
            continue

        add_markdown_paragraph(doc, stripped, cover_mode=cover_mode)
        i += 1

    doc.core_properties.title = "Yapay Zeka Destekli Hibrit Uçuş Öncesi Risk Değerlendirme ve Brifing Asistanı"
    doc.core_properties.subject = "Bitirme Tezi Kitapçığı"
    doc.core_properties.author = "Mete"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
