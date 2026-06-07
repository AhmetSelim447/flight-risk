from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Yapay_Zeka_Destekli_Ucus_Risk_Degerlendirme_Bitirme_Tezi.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def set_table_width(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Yapay Zeka Destekli Uçuş Risk Değerlendirme ve Karar Destek Sistemi")


def para(doc: Document, text: str, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.add_run(text)
    return p


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_kv_table(doc: Document, rows) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [Inches(2.35), Inches(4.0)])
    for i, (key, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_shading(c0, "F2F4F7")
        set_cell_text(c0, key, True)
        set_cell_text(c1, value)


def add_matrix_table(doc: Document, headers, rows, widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        set_table_width(table, widths)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")
    p.add_run(text)


def add_cover(doc: Document) -> None:
    for text, size, bold in [
        ("T.C. FIRAT ÜNİVERSİTESİ", 16, True),
        ("MÜHENDİSLİK FAKÜLTESİ", 14, True),
        ("YAZILIM MÜHENDİSLİĞİ BÖLÜMÜ", 14, True),
    ]:
        p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.runs[0].font.size = Pt(size)
        p.runs[0].bold = bold
    doc.add_paragraph()
    title = para(
        doc,
        "YAPAY ZEKA DESTEKLİ NOTAM VE METAR/TAF ANALİZİ İLE UÇUŞ RİSK DEĞERLENDİRME VE KARAR DESTEK SİSTEMİ",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    title.runs[0].font.size = Pt(17)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph()
    p = para(doc, "BİTİRME TEZİ", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True
    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("Yılı / Dönemi", "2025-2026 Ders Yılı / Bahar Dönemi"),
            ("Öğrenci No", "220290007 - 220290002 - 220290009"),
            ("Ad Soyad", "Mete Han YILMAZ - Ahmet Selim AYTAÇ - Emre NABİKOĞLU"),
            ("Bitirme Tez Danışmanı", "Prof. Dr. Bilal ALATAŞ"),
            ("Proje Başlığı", "Yapay Zeka Destekli NOTAM ve METAR/TAF Analizi ile Uçuş Risk Değerlendirme ve Karar Destek Sistemi"),
            ("Proje Türü", "Açıklanabilir hibrit yapay zeka tabanlı uçuş öncesi karar destek prototipi"),
        ],
    )
    doc.add_paragraph()
    p = para(doc, "Elazığ, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(12)
    add_page_break(doc)


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "Özet", 1)
    for text in [
        "Bu bitirme tezinde, uçuş öncesi operasyonel brifing sürecini desteklemek amacıyla geliştirilen hibrit yapay zeka tabanlı bir karar destek sistemi sunulmaktadır. Sistem; METAR, TAF, NOTAM, pist bilgisi, rüzgar bileşenleri ve havalimanı metaverisini tek bir akışta birleştirerek kullanıcının kalkış ve varış meydanı için anlaşılır bir risk brifingi almasını sağlar.",
        "Çalışmanın temel yaklaşımı, tek başına büyük dil modeline dayanan serbest bir karar mekanizması kurmak yerine, kural tabanlı risk motoru, METAR geçmişinden eğitilmiş operasyonel proxy risk modeli ve NOTAM semantik etki sınıflandırmasını birlikte çalıştırmaktır. Bu nedenle sistemin ürettiği sonuçlar açıklanabilir bileşenlere ayrılır: hava modeli skoru, kural skoru, NOTAM etki skoru, final risk bandı ve güven notu.",
        "Geliştirilen prototip pilot, dispeçer, AIS/AIM, ATC veya resmi operasyonel otoritenin yerine geçmez. Çalışmanın çıktısı, akademik ve prototip düzeyinde bir uçuş öncesi risk destek asistanıdır. Uygulama, özellikle Türkiye LT* meydanları için METAR geçmişinden oluşturulan veri seti ve sentetik/opsiyonel canlı NOTAM sağlayıcıları ile test edilmiştir.",
    ]:
        para(doc, text)
    add_callout(
        doc,
        "Anahtar kelimeler",
        "METAR, TAF, NOTAM, açıklanabilir yapay zeka, hibrit risk modeli, uçuş brifingi, karar destek sistemi, operasyonel proxy risk.",
    )

    add_heading(doc, "Abstract", 1)
    for text in [
        "This thesis presents a hybrid artificial intelligence assisted pre-flight briefing and operational risk-support system. The system combines METAR, TAF, NOTAM, runway information, wind components and airport metadata in a single workflow and produces an explainable briefing for departure and arrival airports.",
        "The proposed architecture does not rely on an unconstrained large language model to decide flight safety. Instead, it combines a deterministic rule engine, a METAR-based operational proxy-risk model and a semantic NOTAM impact layer. The final score is therefore traceable through weather, wind, NOTAM and confidence components.",
        "The prototype is not an operational authority or certified aviation decision system. It is a research and demonstration system designed to show how explainable hybrid AI can support pre-flight risk briefing while preserving transparency, fallback behavior and explicit limitations.",
    ]:
        para(doc, text)
    add_callout(
        doc,
        "Keywords",
        "METAR, TAF, NOTAM, explainable artificial intelligence, hybrid risk model, flight briefing, decision support system, operational proxy risk.",
    )
    add_page_break(doc)

    add_heading(doc, "İçindekiler", 1)
    contents = [
        "1. Giriş",
        "2. Literatür ve Kavramsal Arka Plan",
        "3. Gereksinim Analizi",
        "4. Veri Kaynakları ve Veri Stratejisi",
        "5. Sistem Mimarisi",
        "6. Hibrit Yapay Zeka Risk Modeli",
        "7. NOTAM Analizi ve Sentetik NOTAM Modeli",
        "8. Kullanıcı Arayüzü ve Brifing Deneyimi",
        "9. Test, Doğrulama ve Bulgular",
        "10. Sonuç ve Gelecek Çalışmalar",
        "Kaynakça",
        "Ekler",
    ]
    for item in contents:
        bullet(doc, item)
    add_page_break(doc)


def add_intro(doc: Document) -> None:
    add_heading(doc, "1. Giriş", 1)
    paragraphs = [
        "Havacılık operasyonlarında uçuş öncesi hazırlık, farklı veri kaynaklarının doğru zamanda ve doğru bağlamda yorumlanmasını gerektirir. Bir uçuş planlanırken meteorolojik gözlemler, tahminler, NOTAM bilgileri, pist kullanılabilirliği, rüzgar bileşenleri ve alternatif meydan seçenekleri ayrı ayrı kontrol edilir. Bu kaynakların her biri operasyonel açıdan önemlidir; ancak farklı ekranlarda, farklı formatlarda ve farklı ayrıntı seviyelerinde sunulduklarında kullanıcı için bilişsel yük artar.",
        "Bu çalışmanın çıkış noktası, uçuş öncesi brifing sürecinde dağınık verilerin tek bir karar destek arayüzünde toplanmasıdır. Amaç, kullanıcıya uçuşu onaylayan veya iptal eden bir sistem sunmak değildir. Amaç, mevcut veriye göre hangi konuların dikkat gerektirdiğini açıklanabilir biçimde göstermek, kritik verileri sadeleştirmek ve teknik ayrıntıyı gerektiğinde erişilebilir bırakmaktır.",
        "Proje kapsamında geliştirilen flight-risk uygulaması, kalkış ve varış meydanı seçildiğinde METAR/TAF verilerini canlı sağlayıcı zincirinden alır, NOTAM bilgisini mevcut sağlayıcıya veya sentetik olay motoruna göre üretir, pist ve rüzgar hesaplarını yapar, kural tabanlı risk skorunu çıkarır, AI servisinden hibrit risk tahmini ister ve kullanıcıya tek bir brifing ekranı sunar.",
        "Sistemin en önemli tasarım kararı, yapay zekanın karar verici değil açıklayıcı ve destekleyici rol üstlenmesidir. Büyük dil modeli veya LLM benzeri raporlama katmanı final skoru serbestçe üretmez. Bunun yerine METAR modelinin, kural motorunun ve NOTAM semantik analizinin çıktısını kullanarak anlaşılır bir açıklama üretir. Bu yaklaşım, özellikle havacılık gibi yüksek hassasiyetli alanlarda şeffaflık ihtiyacını karşılamak için tercih edilmiştir.",
    ]
    for text in paragraphs:
        para(doc, text)
    add_heading(doc, "1.1 Problem Tanımı", 2)
    for text in [
        "Klasik uçuş öncesi hazırlıkta METAR ve TAF meteorolojik durumu gösterirken, NOTAM operasyonel kısıtları ve meydan/prosedür durumlarını bildirir. Bu veriler tek başına okunabilir olsa da, uçuş planlama bağlamında birlikte yorumlanmaları gerekir. Örneğin görüş iyi olsa bile pist kapanışı veya ILS arızası varış planını etkileyebilir. Benzer şekilde NOTAM tarafı sakin olsa bile düşük tavan, güçlü yan rüzgar veya kötüleşen TAF eğilimi ek kontrol gerektirebilir.",
        "Sorun yalnızca veriye erişim değildir. Asıl sorun, verinin anlamlı risk başlıklarına ayrılmasıdır. Kullanıcı çoğu zaman 'skor neden 70?', 'hangi NOTAM kritik?', 'hangi veri eksik?', 'bu sonuç güvenilir mi?' gibi sorulara hızlı cevap arar. Bu nedenle sistemin yalnızca sayı üretmesi yeterli değildir; sayının arkasındaki kategori, neden, sınırlılık ve güven seviyesi açıkça gösterilmelidir.",
    ]:
        para(doc, text)
    add_heading(doc, "1.2 Çalışmanın Amacı ve Katkıları", 2)
    for item in [
        "METAR, TAF, NOTAM ve pist/rüzgar bilgisini tek bir brifing orkestrasyonunda birleştirmek.",
        "Kural tabanlı risk motoru ile açıklanabilir ve deterministik bir taban karar desteği sağlamak.",
        "Türkiye LT* METAR geçmişinden eğitilmiş operasyonel proxy risk modeli ile meteorolojik risk bileşenini hesaplamak.",
        "NOTAM metinlerini ve sentetik NOTAM olaylarını kategori, şiddet, etki alanı, pist/prosedür ve skor açısından açıklamak.",
        "LLM benzeri raporlama katmanını final karar verici olarak değil, yalnızca açıklayıcı rapor üretici olarak konumlandırmak.",
        "Kullanıcı arayüzünde sade ilk ekran ve teknik detay paneli ayrımı yaparak tez demosunda anlaşılabilirliği artırmak.",
    ]:
        bullet(doc, item)
    add_page_break(doc)


def add_background(doc: Document) -> None:
    add_heading(doc, "2. Literatür ve Kavramsal Arka Plan", 1)
    sections = [
        (
            "2.1 METAR ve TAF Verisinin Operasyonel Önemi",
            [
                "METAR, meydandaki gözlenen meteorolojik durumu; TAF ise belirli bir zaman aralığı için terminal sahasındaki beklenen koşulları ifade eder. Uçuş öncesi brifingde görüş, rüzgar, bulut tavanı, yağış, gök gürültülü hadise, buzlanma ve trend bilgileri kritik göstergeler olarak değerlendirilir.",
                "AviationWeather Data API, METAR ve TAF ürünleri için dünya çapında kapsama ve JSON, XML, CSV, GeoJSON gibi makine tarafından işlenebilir formatlar sunar. Bu nedenle projede scraping yerine resmi API tabanlı erişim tercih edilmiştir. API ayrıca oran sınırlaması, hata kodları ve cache dosyaları gibi üretim ortamı açısından önemli kuralları da tanımlar.",
                "TAF, geleceğe dönük karar desteği için önemlidir; ancak bu projede TAF henüz ayrı bir ML modelinin hedefi değildir. Bunun yerine TAF canlı brifingde gösterilir, snapshot olarak toplanabilir ve kural/AI açıklama katmanında kötüleşme sinyali olarak değerlendirilir.",
            ],
        ),
        (
            "2.2 NOTAM Verisi ve Anlamsal Zorluk",
            [
                "NOTAM; pist kapanışı, seyrüsefer yardımcısı arızası, ışıklandırma çalışması, apron/taksi yolu kısıtı, hava sahası faaliyeti veya çalışma saati değişikliği gibi uçuş planlamasını etkileyebilecek operasyonel bilgileri taşır. NOTAM metinleri kısa, kodlu ve bağlama bağımlı olduğundan kullanıcı açısından zor yorumlanabilir.",
                "Bu çalışmada NOTAM yalnızca ham metin olarak gösterilmez. Her NOTAM olayının kategori, şiddet, kritik durum, etki alanı, etkilenen pist/prosedür, geçerlilik ve kısa gerekçe bilgisine dönüştürülmesi hedeflenmiştir. Böylece kullanıcı 'kritik' etiketinin hangi operasyona dokunduğunu görebilir.",
                "Canlı NOTAM erişimi için Laminar Data Hub ve SkyLink gibi ticari API seçenekleri, resmi uzun vadeli hedef olarak ise EUROCONTROL EAD ve DHMİ/EAD hattı değerlendirilmiştir. Ancak API anahtarı veya resmi erişim doğrulanmadığında sistem sentetik NOTAM üretir ve bunu açıkça demo/test verisi olarak işaretler.",
            ],
        ),
        (
            "2.3 Açıklanabilir Yapay Zeka ve Karar Destek",
            [
                "Havacılık gibi emniyet hassasiyeti yüksek alanlarda yapay zekanın açıklanabilir olması önemlidir. Kullanıcıya yalnızca final skor verilirse sistemin güvenilirliği ve denetlenebilirliği zayıflar. Bu nedenle risk skoru kategori, bileşen, veri eksikliği ve güven notu ile birlikte sunulmalıdır.",
                "Projede hibrit yaklaşım seçilmiştir. Kural tabanlı motor deterministik ve denetlenebilir bir taban sağlar. ML modeli geçmiş METAR verisinden operasyonel hava riskini tahmin eder. NOTAM semantik katmanı operasyonel kısıtları sayısallaştırır. Raporlama katmanı ise bu çıktıları doğal dile dönüştürür.",
                "Bu yapı, LLM'in kontrolsüz şekilde karar üretmesini engeller. LLM veya LLM-benzeri katman final skoru serbestçe belirlemez; yalnızca sınırlı yorum ve açıklama üretir. Böylece sistemin davranışı hem kullanıcıya hem de jüriye açıklanabilir kalır.",
            ],
        ),
    ]
    for heading, paras in sections:
        add_heading(doc, heading, 2)
        for text in paras:
            para(doc, text)
    add_callout(
        doc,
        "Tasarım ilkesi",
        "Sistem 'uçabilir/uçamaz' kararı vermez; hangi verinin yeniden kontrol edilmesi gerektiğini ve riskin hangi başlıktan kaynaklandığını gösterir.",
    )
    add_page_break(doc)


def add_requirements(doc: Document) -> None:
    add_heading(doc, "3. Gereksinim Analizi", 1)
    para(doc, "Sistem gereksinimleri belirlenirken iki ana hedef dikkate alınmıştır: demo sırasında hızlı anlaşılabilirlik ve teknik olarak hibrit AI mimarisini kanıtlayacak derinlik. Bu nedenle uygulamada ilk ekran sade tutulmuş, teknik ayrıntılar ise genişletilebilir panellerde korunmuştur.")
    add_heading(doc, "3.1 Fonksiyonel Gereksinimler", 2)
    for item in [
        "Kullanıcı kalkış ve varış meydanını ICAO, IATA veya meydan adına göre seçebilmelidir.",
        "Sistem DEP/ARR için METAR ve TAF verisini provider zinciri üzerinden çekebilmelidir.",
        "Sistem NOTAM verisini canlı sağlayıcı varsa oradan, yoksa deterministik sentetik sağlayıcıdan alabilmelidir.",
        "Risk skoru düşük, orta ve yüksek bandında gösterilmeli; skorun nedeni basit maddelerle açıklanmalıdır.",
        "Her kritik NOTAM için kısa operasyonel gerekçe üretilmelidir.",
        "Uçuş risk raporu tablosunda görüş, tavan, yan rüzgar, TAF eğilimi, kritik NOTAM, hava modeli ve model güveni satırları görünmelidir.",
        "Teknik detaylar istendiğinde ML skoru, rule skoru, NOTAM semantik skoru, guardrail ve model versiyonu görülebilmelidir.",
        "PDF brifing çıktısı alınabilmeli ve demo için rota haritada gösterilebilmelidir.",
    ]:
        bullet(doc, item)
    add_heading(doc, "3.2 Fonksiyonel Olmayan Gereksinimler", 2)
    for item in [
        "AI veya dış veri sağlayıcı hatasında sistem tamamen çökmeden kural tabanlı fallback ile çalışmalıdır.",
        "Sentetik veri kullanıcıya açıkça demo/test olarak sunulmalıdır.",
        "Veri kaynağı, fallback kullanımı ve güncellik bilgisi kullanıcıya gösterilmelidir.",
        "UI pilot veya operasyon personeli için okunabilir olmalı; ham metinler varsayılan olarak kapalı kalmalıdır.",
        "Modelin kaza riski değil operasyonel proxy risk modeli olduğu açıkça belirtilmelidir.",
        "Proje mevcut monorepo mimarisi içinde geliştirilmeli, gereksiz mimari kırılımlar yapılmamalıdır.",
    ]:
        bullet(doc, item)
    add_heading(doc, "3.3 Başarı Ölçütleri", 2)
    add_matrix_table(
        doc,
        ["Ölçüt", "Beklenen Sonuç", "Projedeki Karşılığı"],
        [
            ["Anlaşılabilirlik", "Kullanıcı ilk ekranda risk seviyesi ve nedeni görür.", "Karar özeti ve uçuş risk raporu tablosu"],
            ["Açıklanabilirlik", "Skor bileşenleri ve gerekçeler gösterilir.", "ML/rule/NOTAM detay paneli"],
            ["Dayanıklılık", "AI servisi veya provider hatasında sistem çalışır.", "Fallback risk ve simulated NOTAM"],
            ["Akademik kanıt", "Model verisi, metrikleri ve sınırlılıkları belgelenir.", "Kalibrasyon sayfası ve evaluation.json"],
        ],
        [Inches(1.7), Inches(2.3), Inches(2.4)],
    )
    add_page_break(doc)


def add_data_strategy(doc: Document) -> None:
    add_heading(doc, "4. Veri Kaynakları ve Veri Stratejisi", 1)
    para(doc, "Veri stratejisi, canlı brifing ihtiyacı ile model eğitimi ihtiyacını ayrı ele alır. Canlı METAR/TAF verisi kısa vadeli operasyonel brifing için kullanılırken, geçmiş METAR arşivi model eğitimi için kullanılmıştır. NOTAM tarafında canlı resmi erişim doğrulanmadığı için sistem sentetik sağlayıcıyla çalışabilir ve bu durum kullanıcıya açıkça gösterilir.")
    add_heading(doc, "4.1 Canlı METAR/TAF Sağlayıcı Zinciri", 2)
    for text in [
        "Sistemde METAR/TAF sağlayıcı zinciri 'auto' modunda çalışır. İlk kaynak AviationWeather Data API'dir. API, METAR ve TAF ürünlerini dünya çapında ve makine tarafından işlenebilir formatlarda sunar. API'nin oran sınırlamaları ve 204/429/500 gibi durum kodları dikkate alınarak fallback yapısı kurulmuştur.",
        "CheckWX ve AVWX token varsa fallback sağlayıcısı olarak kullanılabilir. Bu servisler yoksa veya başarısız olursa son fallback olarak NOAA text endpoint değerlendirilir. Bu zincir, tek bir sağlayıcıya bağlı kalmadan brifing üretimini sürdürmeyi amaçlar.",
    ]:
        para(doc, text)
    add_heading(doc, "4.2 Tarihsel METAR Veri Seti", 2)
    for text in [
        "Model eğitimi için Iowa State University Iowa Environmental Mesonet ASOS/METAR arşivi kullanılmıştır. Bu arşiv Türkiye ASOS ağı da dahil olmak üzere birçok ülke ve istasyon için geçmiş METAR/ASOS verisini indirebilme olanağı sağlar.",
        "Projede 'turkey' istasyon seçimi, bilinen LT* meydanlarının otomatik seçilmesini sağlar. Eğitim sürecinde 2023-2026 aralığında milyonlarca satırlık METAR kaydı işlenmiş, özellikler çıkarılmış ve operasyonel proxy risk etiketleri üretilmiştir.",
    ]:
        para(doc, text)
    add_heading(doc, "4.3 Meydan ve Pist Verisi", 2)
    para(doc, "Meydan ve pist metaverisi için OurAirports veri setleri kullanılmıştır. airports.csv ve runways.csv dosyaları havalimanı kimliği, koordinat, tip ve pist bilgisi gibi alanları sağlar. Sistem Türkiye meydanlarını filtreleyerek koordinat ve pist bilgisi bulunan kayıtları brifing ve harita modülünde kullanır.")
    add_heading(doc, "4.4 NOTAM Veri Stratejisi", 2)
    for text in [
        "NOTAM tarafında iki katmanlı strateji uygulanmıştır. Canlı provider anahtarı varsa Laminar veya SkyLink gibi sağlayıcılardan veri alınabilir. Anahtar yoksa veya istek başarısız olursa deterministik sentetik NOTAM motoru devreye girer.",
        "Sentetik NOTAM motoru gerçek operasyonel NOTAM yerine geçmez. Aynı ICAO ve aynı zaman bucket'ı için stabil olaylar üretir. Bu olaylar pist kapanışı, pist yüzeyi, seyrüsefer yardımcısı arızası, ışıklandırma, apron/taksi yolu çalışması, hava sahası faaliyeti ve hava uyarısı gibi kategorilerde oluşturulur.",
    ]:
        para(doc, text)
    add_matrix_table(
        doc,
        ["Veri Türü", "Kaynak", "Kullanım Amacı", "Durum"],
        [
            ["METAR/TAF canlı", "AviationWeather Data API", "Brifing ve güncel hava özeti", "Aktif"],
            ["METAR geçmiş", "Iowa Mesonet ASOS/METAR", "ML veri seti ve eğitim", "Aktif"],
            ["Meydan/pist", "OurAirports airports.csv/runways.csv", "Harita, pist ve alternate", "Aktif"],
            ["NOTAM canlı", "Laminar/SkyLink/EAD hedefi", "Gerçek NOTAM entegrasyonu", "Anahtar ile doğrulanacak"],
            ["NOTAM sentetik", "Deterministik event engine", "Demo/test ve pipeline doğrulama", "Aktif"],
        ],
        [Inches(1.3), Inches(1.9), Inches(2.2), Inches(1.1)],
    )
    add_page_break(doc)


def add_architecture(doc: Document) -> None:
    add_heading(doc, "5. Sistem Mimarisi", 1)
    para(doc, "Proje monorepo yapısında geliştirilmiştir. Bu yapı web arayüzü, Express API, ortak tipler, AI servisi ve veri/model araçlarının aynı depo içinde yönetilmesini sağlar. Mimari, mevcut modülleri kırmadan genişletme ilkesine göre tasarlanmıştır.")
    add_heading(doc, "5.1 Genel Akış", 2)
    add_matrix_table(
        doc,
        ["Katman", "Sorumluluk", "Örnek Bileşenler"],
        [
            ["Web UI", "Kullanıcı girişi, brifing ekranı, harita, kalibrasyon", "BriefPanel, SearchBar, HomeDashboard, MapPage"],
            ["Express API", "Brifing orkestrasyonu ve dış servis çağrıları", "/brief, /model/status, /feedback, /brief/logs"],
            ["AI Service", "NOTAM parse, risk predict, briefing report", "/ai/notam/parse, /ai/risk/predict, /ai/brief/report"],
            ["ML Pipeline", "Veri indirme, dataset kurma, eğitim ve validasyon", "tools/ml_pipeline.py"],
            ["Data/Logs", "Feedback ve brifing sorgu kayıtları", "brief_queries.jsonl, brief_feedback.jsonl"],
        ],
        [Inches(1.4), Inches(2.6), Inches(2.4)],
    )
    add_heading(doc, "5.2 Brifing Orkestrasyonu", 2)
    for text in [
        "Kullanıcı kalkış ve varış meydanını seçtiğinde web arayüzü Express API'ye /brief isteği gönderir. API önce meydan ve pist bilgisini çözer. Ardından METAR/TAF provider zincirini çalıştırır, NOTAM sağlayıcısını çağırır ve kural tabanlı risk motorundan baseline skor alır.",
        "API, AI servisinden NOTAM semantik analizi ve risk tahmini ister. AI servisi mevcut risk_model.json dosyası varsa eğitilmiş modeli yükler; yoksa fallback mantığıyla çalışır. Sonuçta API eski response alanlarını bozmadan yeni risk.ml ve aiReport alanlarını opsiyonel olarak ekler.",
        "Bu mimari, AI servisinin çökmesi durumunda brifingin tamamen bozulmasını engeller. Böyle bir durumda kural motoru ve mevcut verilerle sonuç üretilir; confidence düşer ve kullanıcıya sınırlılık gösterilir.",
    ]:
        para(doc, text)
    add_heading(doc, "5.3 Modül Bazlı Geliştirme", 2)
    modules = [
        (
            "Modül 1: Veri Kaynaklarının Analizi ve Genişletilmesi",
            "METAR/TAF için resmi API zinciri kuruldu; metar-taf.com scraping dışlandı; tarihsel METAR için Iowa Mesonet seçildi.",
            "Bu modülde en kritik karar, kırılgan web scraping yerine açık dokümantasyonu bulunan sağlayıcı zinciri kullanmak olmuştur. Böylece sistem hem Türkiye LT* meydanlarını hem de farklı sağlayıcı fallback senaryolarını aynı parser mantığıyla işleyebilecek hale gelmiştir.",
        ),
        (
            "Modül 2: Backend API Geliştirme",
            "Express API brifing orkestratörü olarak korundu; /brief, /model/status, /feedback ve log endpointleri eklendi.",
            "Backend tarafında amaç yalnızca veri taşımak değil, farklı servislerin sonuçlarını tutarlı bir brifing cevabında birleştirmektir. API, AI servisinin ulaşılamadığı durumda bile eski response alanlarını koruyarak uygulamanın çalışmasını sürdürür.",
        ),
        (
            "Modül 3: Harita Tabanlı Görselleştirme",
            "DEP/ARR, rota çizgisi, alternate meydanlar, pist yönü ve rüzgar bilgisi haritada gösterildi.",
            "Harita katmanı, sayısal brifingin görsel karşılığıdır. Kullanıcı rota, varış çevresindeki alternate meydanlar ve rüzgar/pist ilişkisini aynı ekranda görerek riskin mekansal bağlamını daha hızlı anlayabilir.",
        ),
        (
            "Modül 4: Risk Değerlendirme Mantığı",
            "Kural motoru, ML modeli, NOTAM semantik skoru ve guardrail yaklaşımı hibrit final skorda birleştirildi.",
            "Bu modül sistemin karar destek çekirdeğidir. Model çıktısı, deterministic guardrail ve NOTAM etki skoru birlikte kullanıldığı için final skor hem veriye dayalı hem de açıklanabilir kalır.",
        ),
        (
            "Modül 5: Alternate Meydan Önerisi",
            "Alternate önerileri ARR çevresinde, pist uzunluğu, mesafe, NOTAM, rüzgar ve hava uygunluğuna göre sıralandı.",
            "Geliştirme sürecinde alternate önerilerinin kalkış çevresine göre değil, varış meydanı problemine göre sıralanması gerektiği netleştirilmiştir. Bu düzeltme, öneri sistemini operasyonel bağlama daha yakın hale getirmiştir.",
        ),
        (
            "Modül 6: Frontend ve BriefPanel",
            "İlk ekran sadeleştirildi; karar özeti, basit gerekçe ve uçuş risk raporu tablosu öne çıkarıldı.",
            "Arayüzdeki temel iyileştirme, teknik derinliği kaybetmeden ilk bakıştaki karmaşıklığı azaltmaktır. Kullanıcı önce risk bandını ve ana gerekçeyi görür; ML formülü, guardrail ve ham METAR/TAF gibi ayrıntılar gerektiğinde açılır.",
        ),
        (
            "Modül 7: PDF Brifing",
            "AI değerlendirme özeti, hibrit skor, model versiyonu ve METAR guardrail açıklamaları PDF çıktısına eklendi.",
            "PDF çıktısı, uygulama ekranındaki brifingin taşınabilir rapor karşılığıdır. Bu çıktı özellikle tez demosunda sistemin yalnızca ekranda çalışan bir prototip değil, rapor üretebilen bütünleşik bir karar destek aracı olduğunu gösterir.",
        ),
    ]
    for title, desc, detail in modules:
        add_heading(doc, title, 3)
        para(doc, desc)
        para(doc, detail)
    add_page_break(doc)


def add_hybrid_ai(doc: Document) -> None:
    add_heading(doc, "6. Hibrit Yapay Zeka Risk Modeli", 1)
    para(doc, "Bu çalışmanın ayırt edici yönü, yapay zekanın tek parça ve kapalı kutu bir karar mekanizması olarak kullanılmamasıdır. Sistem, farklı risk bileşenlerini ayrı hesaplar ve final skoru ağırlıklı ensemble ile üretir. Böylece riskin hangi başlıktan geldiği kullanıcıya açıklanabilir.")
    add_heading(doc, "6.1 Risk Bileşenleri", 2)
    add_matrix_table(
        doc,
        ["Bileşen", "Ağırlık / Rol", "Açıklama"],
        [
            ["mlScore", "%65", "METAR geçmişinden eğitilmiş operasyonel proxy hava risk modeli"],
            ["ruleScore", "%25", "Görüş, tavan, rüzgar, NOTAM ve temel kurallardan gelen açıklanabilir skor"],
            ["notamSemanticScore", "%10", "NOTAM kategori, şiddet ve kritik etki puanı"],
            ["confidence", "Ayrı gösterge", "Veri eksikliği, fallback ve domain dışı durumları gösterir"],
        ],
        [Inches(1.5), Inches(1.3), Inches(3.5)],
    )
    para(doc, "Final skor formülü şu şekilde uygulanır: finalScore = 0.65 * mlScore + 0.25 * ruleScore + 0.10 * notamSemanticScore. Bu skor 0-39 düşük, 40-69 orta ve 70-100 yüksek risk bandına dönüştürülür.")
    add_heading(doc, "6.2 Model Etiketleme Yaklaşımı", 2)
    for text in [
        "Modelin hedefi gerçek kaza riski değildir. Havacılık kaza verisi çok seyrek, bağlama bağımlı ve etik/hukuki açıdan dikkatli ele alınması gereken bir veri türüdür. Bu nedenle ilk model operasyonel proxy risk etiketleriyle eğitilmiştir.",
        "risk_level=1 düşük görüş, düşük tavan, rüzgar/gust, yağış, sis veya benzeri dikkat gerektiren METAR göstergelerini temsil eder. risk_level=2 ise daha yüksek riskli operasyonel meteorolojik sinyalleri temsil eder. Bu etiketler, sistemi 'emniyet garantisi' değil, 'brifingde dikkat edilmesi gereken başlıkları belirleyen destek aracı' yapar.",
        "Proxy etiket yaklaşımının avantajı, çok daha geniş ve güncel METAR geçmişiyle çalışabilmesidir. Dezavantajı ise etiketlerin gerçek operasyon sonucu veya kaza/olay verisiyle doğrudan eşleşmemesidir. Bu sınırlılık hem UI'da hem de bu tezde açıkça belirtilmiştir.",
    ]:
        para(doc, text)
    add_heading(doc, "6.3 Eğitim ve Validasyon Sonuçları", 2)
    add_matrix_table(
        doc,
        ["Metrik", "Değer"],
        [
            ["Eğitim satırı", "2.024.185"],
            ["Hedef", "risk_level, üç sınıflı operasyonel proxy etiket"],
            ["Normal / Dikkat / Yüksek", "1.843.348 / 85.509 / 95.328"],
            ["Pozitif satır", "180.837"],
            ["ROC AUC", "0.993052161307856"],
            ["Time validation ROC AUC", "0.9948543313889386"],
            ["Airport holdout ROC AUC", "0.9911601733517109"],
            ["Guardrail false negative", "0"],
        ],
        [Inches(2.8), Inches(3.3)],
    )
    add_heading(doc, "6.4 Guardrail Mantığı", 2)
    for text in [
        "Model çıktısı tek başına final kullanıcı riskini belirlemez. Görüş, RVR, tavan, rüzgar/gust, TS/freezing, sis ve yağış gibi açık operasyonel eşikler deterministic guardrail olarak uygulanır. Bu yaklaşım, modelin aşırı iyimser olduğu durumlarda alt risk tabanı oluşturur.",
        "High floor durumunda hava skoru en az 75, final skor ise en az 70 yapılır. Caution floor durumunda hava skoru en az 40, final skor ise en az 40 yapılır. Böylece düşük görüş veya düşük tavan gibi açık durumlar modelin olasılık çıktısına tamamen bırakılmaz.",
        "Guardrail, açıklanabilirlik açısından da önemlidir. UI'da kullanıcı 'hava modeli 75/100' gördüğünde bunun yalnızca istatistiksel model sonucu değil, aynı zamanda belirli eşiklere dayanan koruyucu katmanla birlikte oluştuğunu teknik detayda görebilir.",
    ]:
        para(doc, text)
    add_page_break(doc)


def add_notam_section(doc: Document) -> None:
    add_heading(doc, "7. NOTAM Analizi ve Sentetik NOTAM Modeli", 1)
    para(doc, "NOTAM modülü, projenin hem veri erişimi hem de açıklanabilirlik açısından en önemli parçalarından biridir. NOTAM metni çoğu zaman kısa ve teknik olduğundan doğrudan kullanıcıya verildiğinde risk sebebi net anlaşılmayabilir. Bu nedenle sistem, NOTAM'ı yapılandırılmış olaya dönüştürür.")
    add_heading(doc, "7.1 Sentetik Event Engine", 2)
    for text in [
        "Canlı NOTAM anahtarı olmadığı durumda sistem deterministik sentetik NOTAM sağlayıcısını kullanır. Bu sağlayıcı aynı ICAO ve aynı zaman bucket'ı için aynı olayları üretir. Böylece demo, test ve görsel doğrulama senaryoları kararlı hale gelir.",
        "Olay motoru; meydan profili, pist sayısı, pist uzunluğu, bölge etiketi, UTC saat, sezon ve seed bucket gibi girdilerden kategori ve şiddet seçer. Üretilen olay runway_closure, runway_surface, nav_outage, lighting_maintenance, ops_hours, apron_works, taxiway_works, airspace_activity veya weather_advisory kategorilerinden biri olabilir.",
        "Sentetik olayların her biri score, severity, critical, impacts, validFrom, validTo, affectedRunway ve reason alanlarıyla temsil edilir. LLM benzeri metin katmanı bu yapıyı NOTAM benzeri İngilizce operasyonel metne çevirebilir; ancak severity veya critical değerlerini değiştiremez.",
    ]:
        para(doc, text)
    add_heading(doc, "7.2 Kritik NOTAM Gerekçesi", 2)
    add_matrix_table(
        doc,
        ["Kategori", "Kısa Gerekçe", "Kullanıcıya Gösterilen Anlam"],
        [
            ["Pist kapalı", "Pist doğrudan kullanılamaz.", "Kalkış/iniş planı ve alternate kontrol edilmeli."],
            ["Pist yüzeyi/frenleme", "Yüzey koşulu performansı etkiler.", "Frenleme ve pist performans hesabı tekrar kontrol edilmeli."],
            ["ILS/PAPI/VOR/GNSS etkisi", "Yaklaşma veya seyrüsefer yardımı etkilenir.", "Yaklaşma minima ve yedek prosedür kontrol edilmeli."],
            ["Işıklandırma bakımı", "Görsel rehberlik azalabilir.", "Gece/düşük görüş operasyonu ayrıca değerlendirilmeli."],
            ["Hava sahası kısıtı", "Rota veya irtifa etkilenebilir.", "ATC ve rota planı tekrar kontrol edilmeli."],
            ["Çalışma saati kısıtı", "Meydan operasyon penceresi değişir.", "Uçuş zamanı ve meydan açıklığı doğrulanmalı."],
        ],
        [Inches(1.8), Inches(2.0), Inches(2.5)],
    )
    add_heading(doc, "7.3 Canlı NOTAM Sağlayıcıları", 2)
    for text in [
        "Laminar Data Hub NOTAM API, kullanıcı anahtarı ile sorgulanabilen global kapsamlı NOTAM veri API'leri sunar. SkyLink gibi ticari sağlayıcılar da NOTAM verisini REST API üzerinden sunabilir. Ancak bu kaynakların gerçek projede kullanılabilmesi için API anahtarı, erişim koşulları, veri lisansı ve Türkiye kapsaması ayrı ayrı doğrulanmalıdır.",
        "Uzun vadeli resmi kaynak olarak EUROCONTROL EAD ve Türkiye için DHMİ/EAD hattı değerlendirilmelidir. Bu tez kapsamında canlı NOTAM sağlayıcı entegrasyonu fallback-ready tasarlanmış, fakat geçerli API anahtarı olmadığı durumda sentetik sağlayıcı açıkça demo/test olarak korunmuştur.",
    ]:
        para(doc, text)
    add_page_break(doc)


def add_ui_section(doc: Document) -> None:
    add_heading(doc, "8. Kullanıcı Arayüzü ve Brifing Deneyimi", 1)
    para(doc, "Frontend tasarımında hedef, teknik ayrıntı ile demo anlaşılabilirliği arasında denge kurmaktır. İlk denemelerde ekranın çok fazla metin, ham METAR/TAF ve teknik skor bileşeni göstermesi kullanıcı açısından yorucu bulunmuştur. Son ürün güncellemesinde ilk ekran sadeleştirilmiş ve teknik detaylar kapalı panellere taşınmıştır.")
    add_heading(doc, "8.1 İlk Ekran Tasarımı", 2)
    for item in [
        "Rota özeti: DEP ve ARR meydanı, aktif pist ve veri kaynağı.",
        "Risk seviyesi: düşük, orta veya yüksek risk bandı.",
        "En önemli gerekçeler: kritik NOTAM, hava modeli, eksik görüş/tavan veya rüzgar etkisi.",
        "Uçuş risk raporu tablosu: her parametre için iyi/izle/risk/eksik durumu.",
    ]:
        bullet(doc, item)
    para(doc, "Bu yapı sayesinde jüri veya kullanıcı ilk bakışta rotayı, risk seviyesini, riskin neden oluştuğunu ve hangi verinin eksik/problemli olduğunu anlayabilir.")
    add_heading(doc, "8.2 Teknik Detayların Ayrılması", 2)
    para(doc, "ML skoru, rule skoru, NOTAM semantik skoru, guardrail, model formülü, model versiyonu, confidence faktörleri ve feedback sistemi ilk ekranda gösterilmez. Bunlar 'Model, NOTAM ve geri bildirimi göster' gibi genişletilebilir panellerde korunur. Böylece sistem hem sade demo hem teknik savunma ihtiyacını karşılar.")
    add_heading(doc, "8.3 METAR/TAF Sunumu", 2)
    for text in [
        "METAR ve TAF bilgisi pilotlar için önemli olsa da ham format herkes için hızlı okunabilir değildir. Bu nedenle sistem özet kartlarda rüzgar, görüş, tavan ve hadise gibi ana alanları sadeleştirir. Ham METAR/TAF metni ve ayrıştırılmış zaman çizelgesi ise detay altında verilir.",
        "TAF tarafında belirgin kötüleşme sinyali yoksa bu bilgi açıkça gösterilir. Görüş veya tavan ayrıştırılamazsa tablo satırı 'eksik' olarak işaretlenir. Bu, verinin yok sayılmasını değil, kullanıcıya veri kalitesinin görünür kılınmasını sağlar.",
    ]:
        para(doc, text)
    add_heading(doc, "8.4 Feedback ve Kalibrasyon", 2)
    para(doc, "Briefing feedback paneli kullanıcının sonucu doğru, fazla muhafazakar, kaçan risk veya yanlış neden olarak işaretlemesine izin verir. Bu veriler şimdilik otomatik eğitime girmemektedir; ancak gelecekte model kalibrasyonu ve eşik ayarı için yerel manuel etiket kaynağı oluşturur.")
    add_page_break(doc)


def add_validation(doc: Document) -> None:
    add_heading(doc, "9. Test, Doğrulama ve Bulgular", 1)
    para(doc, "Doğrulama süreci üç düzeyde ele alınmıştır: veri/model doğrulaması, API entegrasyon doğrulaması ve kullanıcı arayüzü doğrulaması. Model tarafında time split ve airport holdout sonuçları incelenmiştir. API tarafında fallback davranışı, log endpointleri ve provider metadata alanları kontrol edilmiştir. UI tarafında ise risk gerekçelerinin anlaşılabilirliği önceliklendirilmiştir.")
    add_heading(doc, "9.1 Model Validasyonu", 2)
    add_matrix_table(
        doc,
        ["Validasyon", "Satır", "ROC AUC", "False Negative", "False Positive", "Guardrail FN"],
        [
            ["Time validation", "304.920", "0.994854", "384", "10.769", "0"],
            ["Airport holdout", "218.157", "0.991160", "1.373", "7.378", "0"],
        ],
        [Inches(1.6), Inches(1.0), Inches(1.0), Inches(1.2), Inches(1.2), Inches(1.0)],
    )
    for text in [
        "ROC AUC değerlerinin yüksek olması modelin proxy etiketleri iyi ayırabildiğini gösterir. Ancak bu sonuçlar gerçek uçuş emniyeti veya kaza riski tahmini olarak yorumlanmamalıdır. Çünkü etiketler METAR operasyonel göstergelerinden türetilmiştir.",
        "Guardrail sonrası false negative sayısının sıfırlanması, açık meteorolojik risk eşiklerinin kullanıcı tarafında düşük gösterilmesini engelleyen koruyucu mantığın çalıştığını gösterir. Bunun bedeli daha fazla false positive, yani daha muhafazakar sonuç üretme eğilimidir.",
    ]:
        para(doc, text)
    add_heading(doc, "9.2 Entegrasyon Senaryoları", 2)
    for item in [
        "LTFM -> LTAC rotasında METAR/TAF provider bilgisi, sentetik NOTAM, risk skoru ve AI raporu tek ekranda gösterilmiştir.",
        "AI servisi yoksa /brief response eski alanlarını koruyarak kural tabanlı fallback ile çalışacak şekilde tasarlanmıştır.",
        "NOTAM sağlayıcı simulated olduğunda UI bunu demo/test verisi olarak gösterir.",
        "Alternatif meydan önerileri varış meydanı çevresinde sıralanır; kalkış çevresine göre yanlış öneri üretilmemesi hedeflenmiştir.",
        "Kritik NOTAM varsa genel 'kritik NOTAM var' cümlesi yerine pist, yüzey, seyrüsefer veya hava uyarısı gibi somut neden gösterilir.",
    ]:
        bullet(doc, item)
    add_heading(doc, "9.3 Kullanıcı Deneyimi Bulguları", 2)
    para(doc, "Sadeleştirme öncesi arayüzde çok fazla teknik bilgi aynı anda görünmekteydi. Skorlar ve breakdown alanları kullanıcının 'neden bu risk?' sorusunu doğrudan cevaplamıyordu. Son düzenlemede karar özeti, basit gerekçe ve uçuş risk raporu tablosu öne çıkarılmış; ham ve teknik alanlar kapalı panellere taşınmıştır.")
    add_page_break(doc)


def add_results_and_conclusion(doc: Document) -> None:
    add_heading(doc, "10. Sonuç ve Gelecek Çalışmalar", 1)
    for text in [
        "Bu tez kapsamında, uçuş öncesi operasyonel brifing sürecini destekleyen açıklanabilir hibrit yapay zeka tabanlı bir prototip geliştirilmiştir. Sistem, METAR/TAF verisi, NOTAM bilgisi, pist/rüzgar hesapları, ML modeli ve kural tabanlı risk motorunu tek bir brifing akışında birleştirir.",
        "Çalışmanın en önemli katkısı, yapay zekayı karar verici yerine açıklanabilir karar destek bileşeni olarak konumlandırmasıdır. ML modeli yalnızca METAR tabanlı operasyonel proxy risk üretir. NOTAM semantik katmanı operasyonel kısıtları sınıflandırır. LLM benzeri raporlama katmanı ise bu bileşenleri kullanıcı dostu bir açıklamaya dönüştürür.",
        "Uygulama, tez demosu için anlaşılır bir ilk ekran ve teknik savunma için detaylı analiz panelleri sağlar. Böylece jüri hem ürün değerini hem de arka plandaki hibrit AI mimarisini görebilir.",
    ]:
        para(doc, text)
    add_heading(doc, "10.1 Sınırlılıklar", 2)
    for item in [
        "Canlı Türkiye NOTAM entegrasyonu henüz doğrulanmış API anahtarıyla test edilmemiştir.",
        "TAF geçmiş veri seti henüz tamamlanmamıştır; TAF ayrı bir ML modeli olarak eğitilmemiştir.",
        "Model gerçek kaza/olay verisiyle değil, METAR operasyonel proxy etiketleriyle eğitilmiştir.",
        "Feedback verileri toplanmakta ancak otomatik model eğitimine henüz dahil edilmemektedir.",
        "Sistem operasyonel sertifikalı karar sistemi değildir; resmi brifing ve otorite kararının yerine geçmez.",
    ]:
        bullet(doc, item)
    add_heading(doc, "10.2 Gelecek Çalışmalar", 2)
    for item in [
        "DHMİ/EAD/EUROCONTROL veya lisanslı ticari API üzerinden gerçek Türkiye NOTAM entegrasyonunun doğrulanması.",
        "TAF snapshot birikimiyle TAF trend veri seti oluşturulması ve ayrı TAF ML skorunun eklenmesi.",
        "Feedback etiketlerinin model kalibrasyonu ve threshold tuning sürecine dahil edilmesi.",
        "Operasyonel delay/diversion gibi daha güçlü proxy etiketlerin Türkiye veya bölgesel veriyle araştırılması.",
        "Model açıklanabilirliği için SHAP benzeri özellik katkısı gösterimlerinin eklenmesi.",
        "PDF çıktısının kurumsal brifing formatına daha da yaklaştırılması ve dışa aktarım seçeneklerinin artırılması.",
    ]:
        bullet(doc, item)
    add_page_break(doc)


def add_sources(doc: Document) -> None:
    add_heading(doc, "Kaynakça", 1)
    sources = [
        "[1] Aviation Weather Center, AviationWeather Data API. https://aviationweather.gov/data/api/",
        "[2] Iowa State University Iowa Environmental Mesonet, ASOS/AWOS/METAR Data Download. https://mesonet.agron.iastate.edu/request/download.phtml",
        "[3] OurAirports, Open Data: airports.csv and runways.csv. https://ourairports.com/data/",
        "[4] Laminar Data Hub, NOTAM Data APIs v2. https://developer.laminardata.aero/documentation/notamdata/v2",
        "[5] EUROCONTROL, European AIS Database. https://www.eurocontrol.int/service/european-ais-database",
        "[6] FAA SWIFT Portal / SWIM. https://portal.swim.faa.gov/",
        "[7] ICAO, Aeronautical Information Services and NOTAM concepts. https://www.icao.int/airnavigation/AIS",
        "[8] Node.js Documentation. https://nodejs.org/en/docs",
        "[9] React Documentation. https://react.dev/",
        "[10] Vite Documentation. https://vite.dev/",
        "[11] Express Documentation. https://expressjs.com/",
        "[12] FastAPI Documentation. https://fastapi.tiangolo.com/",
        "[13] scikit-learn Documentation. https://scikit-learn.org/",
        "[14] Proje kaynak kodu ve PROJECT_BRAIN.md, flight-risk monorepo, 2026.",
    ]
    for source in sources:
        para(doc, source)
    add_page_break(doc)


def add_appendices(doc: Document) -> None:
    add_heading(doc, "Ek A: API ve Servis Uçları", 1)
    add_matrix_table(
        doc,
        ["Endpoint", "Görev", "Not"],
        [
            ["/brief", "DEP/ARR için tam brifing üretir.", "Eski alanlar korunur; risk.ml ve aiReport opsiyoneldir."],
            ["/model/status", "Model, provider, snapshot ve validasyon durumunu döndürür.", "Kalibrasyon sayfasında kullanılır."],
            ["/feedback", "Kullanıcı brifing geri bildirimi alır.", "Yerel manuel etiket kaynağıdır."],
            ["/brief/logs/latest", "Son brifing sorgusunu döndürür.", "Demo ve debugging için kullanılır."],
            ["/ai/notam/parse", "NOTAM metnini yapılandırılmış etkiye çevirir.", "AI servisinde çalışır."],
            ["/ai/risk/predict", "Hibrit risk tahmini üretir.", "ML/rule/NOTAM skorlarını birleştirir."],
            ["/ai/brief/report", "Açıklanabilir brifing raporu üretir.", "LLM-benzeri rapor katmanıdır."],
        ],
        [Inches(1.8), Inches(2.8), Inches(1.8)],
    )
    add_heading(doc, "Ek B: Veri ve Model Komutları", 1)
    commands = [
        "python tools/ml_pipeline.py download-metar --stations turkey --start 2023-01-01 --end 2026-01-01 --pause 5",
        "python tools/ml_pipeline.py collect-live --stations turkey --kinds taf",
        "npm run data:build-dataset",
        "npm run ml:train",
        "npm run ml:evaluate",
        "python -m py_compile tools/ml_pipeline.py services/nlp/main.py",
        "npm --prefix apps/api run build",
        "npm --prefix apps/web run build",
    ]
    for cmd in commands:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    add_heading(doc, "Ek C: Demo Senaryosu", 1)
    steps = [
        "Uygulama start-flight-risk.bat ile başlatılır.",
        "Kullanıcı DEP olarak LTFM veya LTAC, ARR olarak LTAC veya LTCA seçer.",
        "Brifing Al butonuna basıldığında API METAR/TAF, NOTAM, pist ve model sonuçlarını toplar.",
        "İlk ekranda risk bandı, en önemli gerekçeler ve uçuş risk raporu görülür.",
        "Teknik soru gelirse model/NOTAM/feedback detay paneli açılır.",
        "Harita ve PDF çıktısı üzerinden brifing görsel olarak desteklenir.",
    ]
    for step in steps:
        numbered(doc, step)
    add_heading(doc, "Ek D: Tez Savunması İçin Kısa Mesaj", 1)
    for text in [
        "Bu proje, uçuş öncesi karar destek sürecinde yapay zekanın nasıl güvenli ve açıklanabilir bir yardımcı rol üstlenebileceğini göstermektedir. Sistem pilot yerine karar vermez; farklı kaynaklardan gelen bilgiyi birleştirir, risk başlıklarını kategorize eder ve kullanıcıya hangi konuların tekrar kontrol edilmesi gerektiğini gösterir.",
        "Hibrit yapının farkı, her bileşenin ayrı sorumluluk taşımasıdır. Kural motoru açıklanabilir tabanı sağlar, METAR modeli geçmiş veriden hava riskini öğrenir, NOTAM semantik katmanı operasyonel kısıtları sınıflandırır, LLM-benzeri rapor katmanı ise teknik çıktıları okunabilir hale getirir. Bu nedenle sistem hem teknik olarak savunulabilir hem de kullanıcı açısından anlaşılırdır.",
    ]:
        para(doc, text)


def add_long_discussion(doc: Document) -> None:
    add_heading(doc, "11. Tartışma", 1)
    topics = [
        (
            "11.1 Neden Hibrit Mimari?",
            [
                "Yalnızca kural tabanlı bir sistem, açık eşikleri iyi açıklasa da geçmiş veriden örüntü öğrenemez. Örneğin düşük görüş veya güçlü rüzgar gibi durumlar eşiklerle yakalanabilir; ancak farklı hava koşullarının birlikte oluşturduğu operasyonel baskıyı yalnızca statik kurallarla temsil etmek sınırlı kalır.",
                "Yalnızca ML tabanlı bir sistem ise metrik olarak başarılı olsa bile kullanıcıya 'neden' sorusunun cevabını her zaman doğal biçimde veremez. Bu tezde model çıktısı kararın tamamı değildir; yalnızca METAR tabanlı hava risk bileşenidir. Böylece modelin yetki sınırı teknik olarak daraltılmıştır.",
                "Yalnızca LLM tabanlı bir sistem ise havacılık gibi alanlarda kontrolsüz yorum üretme riski taşır. Bu nedenle LLM-benzeri rapor katmanı, final skor üreticisi değil açıklama üreticisidir. Hibrit mimari her yaklaşımın güçlü yönünü ayrı sorumluluk alanında kullanır.",
            ],
        ),
        (
            "11.2 Proxy Etiketlerin Akademik Değeri",
            [
                "Gerçek kaza veya olay verisiyle çalışmak hem veri erişimi hem de etik açıdan zordur. Ayrıca kaza, tek başına METAR veya NOTAM ile açıklanamayacak kadar çok değişkenli bir sonuçtur. İnsan faktörü, bakım durumu, operasyonel kararlar, uçak tipi ve şirket prosedürleri gibi değişkenler dışarıda kaldığında kaza riski modeli iddiası teknik olarak savunulamaz.",
                "Bu tezde proxy etiketler, sistemin operasyonel brifing riskini öğrenmesi için kullanılmıştır. Düşük görüş, düşük tavan, RVR, rüzgar/gust, sis, yağış, gök gürültülü hadise ve freezing sinyalleri operasyonel dikkat gerektiren meteorolojik göstergeler olarak etiketlenmiştir.",
                "Proxy etiket yaklaşımı sınırlıdır; ancak prototipin veri yoğun şekilde eğitilmesini ve açıklanabilir risk bandı üretmesini sağlar. Bu sınırlılık saklanmamış, UI ve tez metninde açıkça belirtilmiştir. Böylece akademik dürüstlük korunurken uygulanabilir bir ML pipeline oluşturulmuştur.",
            ],
        ),
        (
            "11.3 Sentetik NOTAM'ın Rolü",
            [
                "Sentetik NOTAM gerçek NOTAM değildir; fakat geliştirme ve demo aşamasında kararlı senaryo üretimi sağlar. Canlı NOTAM anahtarı olmadığı durumda pipeline'ın geri kalanını test etmek, UI tasarımını doğrulamak ve NOTAM semantik skorunun davranışını göstermek için değerlidir.",
                "Deterministik üretim sayesinde aynı meydan ve aynı zaman bucket'ı için aynı olaylar oluşur. Bu özellik test senaryolarında önemlidir; çünkü her demo çalıştırmasında farklı veri oluşması, hatanın UI'dan mı yoksa veri değişiminden mi kaynaklandığını anlamayı zorlaştırır.",
                "Bu nedenle sistemde sentetik NOTAM saklanmamış veya gizlenmemiş, tersine açıkça demo/test olarak işaretlenmiştir. Kullanıcı arayüzünde canlı ve sentetik ayrımı korunarak gerçek operasyonel NOTAM algısı yaratılmasının önüne geçilmiştir.",
            ],
        ),
        (
            "11.4 Kullanıcı Deneyimi Açısından Dersler",
            [
                "İlk prototipte çok fazla teknik bilgi aynı anda gösterildiğinde kullanıcı riskin neden oluştuğunu daha zor anlamıştır. ML skoru, rule skoru, guardrail, primary driver, breakdown ve ham NOTAM metni aynı anda görünür olduğunda ana mesaj kaybolmuştur.",
                "Son tasarımda risk seviyesi, basit gerekçeler ve uçuş risk raporu tablosu ilk ekrana alınmıştır. Teknik ayrıntılar tamamen kaldırılmamış, yalnızca kapalı panellere taşınmıştır. Bu yaklaşım tez demosunda anlatılabilirliği artırdığı gibi gerçek kullanıcı deneyimi açısından da daha doğru bir bilgi hiyerarşisi sağlar.",
                "Kritik NOTAM gerekçelerinin madde madde gösterilmesi de önemli bir ders olmuştur. Kullanıcı 'kritik NOTAM: 5' bilgisinden çok, bu beş notamın pist kapanışı mı, seyrüsefer arızası mı, yüzey/frenleme mi olduğunu görmek ister.",
            ],
        ),
        (
            "11.5 Emniyet ve Sorumluluk Sınırı",
            [
                "Sistemin en önemli sınırı, operasyonel otorite yerine geçmemesidir. Bu ifade yalnızca hukuki bir not değil, mimari bir ilkedir. Final skorun adı risk desteğidir; sistem kullanıcıya resmi NOTAM, METAR/TAF, şirket prosedürü ve yetkili karar süreçlerini kontrol etmesi gerektiğini hatırlatır.",
                "Hibrit modelin sonucu emniyet garantisi değildir. Özellikle canlı NOTAM provider doğrulanmadığında veya METAR/TAF eksik olduğunda confidence düşürülür ve eksik veri kullanıcıya gösterilir. Eksik veriyi saklamak yerine görünür yapmak, karar destek sisteminin güvenilirliği açısından temel ilkedir.",
                "Bu sınırın açıkça yazılması, projenin değerini azaltmaz; aksine doğru konumlandırır. Tezin savunulabilir noktası, sistemin resmi karar yerine geçmesi değil, farklı kaynaklardan gelen operasyonel sinyalleri açıklanabilir biçimde bir araya getirmesidir.",
            ],
        ),
    ]
    for heading, paragraphs in topics:
        add_heading(doc, heading, 2)
        for text in paragraphs:
            para(doc, text)
        add_callout(doc, "Değerlendirme", "Bu başlık, sistemin tez savunmasında yalnızca çalışan bir uygulama değil, gerekçeli bir mühendislik kararı olarak anlatılmasını sağlar.")
    add_page_break(doc)


def add_case_studies(doc: Document) -> None:
    add_heading(doc, "12. Kullanım Senaryoları ve Demo Akışı", 1)
    para(doc, "Bu bölümde sistemin tez demosunda nasıl anlatılabileceği ve farklı uçuş öncesi durumlarda hangi karar destek bilgisini ürettiği senaryolar üzerinden açıklanmaktadır. Senaryolar gerçek operasyonel uçuş kararı yerine geçmez; yalnızca prototipin davranışını ve mimari katkısını göstermeyi amaçlar.")

    scenarios = [
        (
            "12.1 Orta Riskli Rota Senaryosu",
            "Kullanıcı LTFM -> LTAC gibi yaygın bir rota seçtiğinde sistem önce canlı METAR/TAF verisini provider zincirinden alır. NOTAM provider simulated ise bu durum veri durumu kartında açıkça gösterilir. Risk seviyesi 40-69 bandında olduğunda arayüz bunu 'Orta Risk' olarak sunar ve uçuşu otomatik engellemediğini, ancak limit, alternate ve güncel veri kontrolü gerektirdiğini belirtir.",
            "Bu senaryoda beklenen çıktı, kullanıcının neden orta risk gördüğünü anlamasıdır. Örneğin NOTAM etkisi yüksek fakat meteorolojik değerler iyi ise sistem ana nedeni NOTAM olarak gösterir. Uçuş risk raporu tablosunda görüş, tavan, yan rüzgar, TAF eğilimi ve kritik NOTAM satırları ayrı ayrı incelenebilir.",
        ),
        (
            "12.2 Yüksek Riskli NOTAM Senaryosu",
            "Varış meydanında pist kapanışı, pist yüzeyi/frenleme problemi veya ILS/PAPI/VOR/GNSS etkisi gibi kritik NOTAM üretildiğinde sistem yalnızca 'kritik NOTAM var' demez. DEP ve ARR ayrımıyla hangi meydanda hangi operasyonel başlığın sorun çıkardığını gösterir.",
            "Bu senaryonun tez açısından önemi, NOTAM semantik katmanının kullanıcı değerini göstermesidir. Ham NOTAM metni detay altında kalırken ilk ekranda pist kapalı, seyrüsefer yardımı etkilenmiş veya hava sahası kısıtı var gibi doğrudan anlaşılabilir maddeler yer alır.",
        ),
        (
            "12.3 Eksik Veri Senaryosu",
            "Bazı meydanlarda METAR veya TAF alanları eksik, gecikmiş veya ayrıştırılamamış olabilir. Sistem bu eksikliği gizlemek yerine risk raporu tablosunda 'eksik' olarak gösterir. Böylece kullanıcı skorun hangi veriye dayanmadığını ve güven seviyesinin neden düştüğünü anlayabilir.",
            "Eksik veri yaklaşımı özellikle havacılıkta önemlidir. Çünkü eksik bilginin normal bilgi gibi davranması hatalı güven oluşturur. Bu projede confidence faktörleri; TAF verisi, rüzgar verisi, ceiling/görüş ayrıştırma durumu ve provider fallback bilgisiyle birlikte değerlendirilir.",
        ),
        (
            "12.4 AI Servisi Kapalı Senaryosu",
            "AI servisi çalışmıyorsa Express API tamamen durmaz. Kural tabanlı risk motoru ve mevcut provider verileriyle brifing üretimi devam eder. Bu durumda risk.ml alanları fallback model versiyonu veya düşük confidence ile işaretlenebilir.",
            "Bu senaryo, sistemin dayanıklılık tasarımını gösterir. AI katmanı değer katar; fakat temel brifing akışı ona tek noktadan bağımlı değildir. Tez savunmasında bu özellik, yüksek riskli alanlarda fail-soft yaklaşımın örneği olarak anlatılabilir.",
        ),
    ]

    for heading, p1, p2 in scenarios:
        add_heading(doc, heading, 2)
        para(doc, p1)
        para(doc, p2)

    add_heading(doc, "12.5 Demo Sırasında Vurgulanacak Noktalar", 2)
    for item in [
        "Sistem pilot yerine karar vermez; karar destek ve brifing asistanıdır.",
        "Yapay zeka tek başına final skoru yazmaz; ML, kural ve NOTAM skorları birleşir.",
        "Sentetik NOTAM varsa açıkça demo/test verisi olarak görünür.",
        "Skorun nedeni ilk ekranda basit maddelerle, teknik ayrıntısı detay panelinde gösterilir.",
        "Model kaza riski değil, METAR tabanlı operasyonel proxy risk modelidir.",
    ]:
        bullet(doc, item)
    add_page_break(doc)


def add_engineering_decisions(doc: Document) -> None:
    add_heading(doc, "13. Mühendislik Kararları ve Risk Yönetimi", 1)
    para(doc, "Proje yalnızca bir arayüz çalışması değil, aynı zamanda veri güvenilirliği, servis dayanıklılığı ve açıklanabilir AI kararları üzerine kurulmuş bir mühendislik çalışmasıdır. Bu bölümde uygulama geliştirilirken alınan ana kararlar ve bu kararların gerekçeleri özetlenmektedir.")

    add_heading(doc, "13.1 Scraping Yerine Provider Zinciri", 2)
    for text in [
        "METAR/TAF verisi için scraping yapılmaması bilinçli bir karardır. metar-taf.com gibi siteler manuel referans için yararlı olabilir; ancak anti-bot davranışı, izin/ToS belirsizliği ve kırılgan HTML yapısı nedeniyle üretim sağlayıcısı olarak uygun değildir.",
        "Provider zinciri yaklaşımı, AviationWeather API'yi birincil kaynak yaparken CheckWX, AVWX ve NOAA text endpoint gibi fallback seçeneklerini korur. Böylece tek sağlayıcı hatası brifing akışını tamamen bozmaz.",
    ]:
        para(doc, text)

    add_heading(doc, "13.2 Teknik Detayı Saklamadan Sadeleştirme", 2)
    for text in [
        "İlk UI denemelerinde teknik açıklama ile kullanıcı açıklaması aynı alanda karışmıştır. Son sürümde bu iki ihtiyaç ayrılmıştır. İlk ekran karar özeti ve risk raporu tablosu verir; teknik panel ise jüri veya geliştirici sorusu geldiğinde açılır.",
        "Bu ayrım, akademik sunum için özellikle önemlidir. Jüri ilk bakışta ürün değerini görür; detay sorulduğunda model versiyonu, score formula, guardrail ve confidence faktörleri gösterilebilir.",
    ]:
        para(doc, text)

    add_heading(doc, "13.3 Log ve Geri Bildirim Tasarımı", 2)
    para(doc, "Sistemde brifing sorgularının loglanması, model davranışını daha sonra incelemek için gereklidir. data/logs/brief_queries.jsonl dosyası ve /brief/logs/latest endpointi, hangi DEP/ARR sorgusunda hangi sonuçların döndüğünü kontrol etmeyi kolaylaştırır. Feedback paneli ise ileride threshold validation ve manuel kalibrasyon için veri toplar.")

    add_heading(doc, "13.4 Sürdürme ve Genişletme Stratejisi", 2)
    para(doc, "Mimari mevcut monorepo içinde korunmuştur. Express API brifing orkestratörü, services/nlp AI katmanı, tools/ml_pipeline.py veri/model pipeline'ı ve apps/web kullanıcı arayüzü belirgin sorumluluklara sahiptir. Bu ayrım, ileride TAF ML modeli veya canlı NOTAM provider doğrulaması eklendiğinde mevcut sistemin kökten değişmeden genişlemesini sağlar.")
    add_matrix_table(
        doc,
        ["Karar", "Gerekçe", "Risk Azaltımı"],
        [
            ["Provider zinciri", "Tek canlı kaynağa bağımlılığı azaltır.", "Fallback ile brifing devam eder."],
            ["Sentetik NOTAM etiketi", "Demo verisinin yanlış anlaşılmasını engeller.", "Operasyonel veri yanılsaması azalır."],
            ["Guardrail", "Açık meteorolojik riskleri modelden bağımsız yakalar.", "Aşırı iyimser skor riski düşer."],
            ["Detay panelleri", "İlk ekranı sade tutar.", "Teknik derinlik korunur."],
            ["Feedback/log", "Sonuçların sonradan denetlenmesini sağlar.", "Kalibrasyon için veri birikir."],
        ],
        [Inches(1.7), Inches(2.3), Inches(2.2)],
    )
    add_page_break(doc)


def main() -> None:
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_intro(doc)
    add_background(doc)
    add_requirements(doc)
    add_data_strategy(doc)
    add_architecture(doc)
    add_hybrid_ai(doc)
    add_notam_section(doc)
    add_ui_section(doc)
    add_validation(doc)
    add_results_and_conclusion(doc)
    add_long_discussion(doc)
    add_sources(doc)
    add_appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
